#!/usr/bin/env python3
"""Document Editor Agent - Applies improvement plans to DOCX files."""

import logging
import os
import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from ..contracts import Plan, ResumeAST, validate_or_raise

logger = logging.getLogger(__name__)


class DocumentEditorAgent:
    """Agent responsible for applying improvement plans to DOCX documents."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def apply_plan(self, docx_path: str, plan: Plan, resume_ast: ResumeAST, structure_tags: Optional[Dict] = None, original_company_lines: Optional[List[str]] = None) -> str:
        """
        Apply improvement plan to DOCX document.
        
        Args:
            docx_path: Path to original DOCX file
            plan: Improvement plan with edits to apply
            resume_ast: Parsed resume structure for context
            
        Returns:
            str: Path to modified DOCX file
            
        Raises:
            ValueError: If plan cannot be applied safely
            FileNotFoundError: If DOCX file doesn't exist
        """
        try:
            if not os.path.exists(docx_path):
                raise FileNotFoundError(f"DOCX file not found: {docx_path}")
            
            # Handle both Plan object and dictionary
            if hasattr(plan, 'edits'):
                edits = plan.edits
            else:
                edits = plan.get('edits', [])
            
            self.logger.info(f"Applying plan with {len(edits)} edits to {docx_path}")
            
            # Load document
            doc = Document(docx_path)
            
            # Create mapping between ResumeAST and document
            if structure_tags and structure_tags.get("role_spans"):
                bullet_mapping, role_mapping = self._build_mappings_from_tags(doc, resume_ast, structure_tags)
                self.logger.info("Using StructureTaggerAgent mappings for roles and bullets")
            else:
                bullet_mapping = self._create_bullet_mapping(doc, resume_ast)
                role_mapping = self._create_role_mapping(doc, resume_ast)
            
            # Deterministic ordering: modify -> insert -> upsert_skill
            def _edit_key(e):
                et = e.get('type') if isinstance(e, dict) else e.type
                order = {'modify_bullet': 0, 'insert_bullet': 1, 'upsert_skill': 2}
                return order.get(et, 99)
            ordered_edits = sorted(edits, key=_edit_key)

            # Apply each edit deterministically
            for i, edit in enumerate(ordered_edits):
                edit_type = edit.get('type') if isinstance(edit, dict) else edit.type
                self.logger.info(f"Applying edit {i+1}/{len(edits)}: {edit_type}")
                self._apply_single_edit(doc, edit, bullet_mapping, role_mapping, resume_ast)
            
            # No cleanup - only apply the specific edits requested
            self.logger.info("Applied all requested edits - no automatic cleanup performed")
            
            # Preserve company lines if requested
            try:
                if original_company_lines:
                    self._ensure_company_lines_preserved(doc, resume_ast, role_mapping, original_company_lines)
            except Exception as e:
                self.logger.warning(f"Company line preservation step failed: {e}")

            # Save modified document
            output_path = self._generate_output_path(docx_path)
            doc.save(output_path)
            
            self.logger.info(f"Successfully applied plan. Modified file saved to: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to apply plan: {e}")
            raise

    def _build_mappings_from_tags(self, doc: Document, resume_ast: ResumeAST, structure_tags: Dict) -> Tuple[Dict[str, int], Dict[str, Tuple[int, int]]]:
        """Build role and bullet mappings using LLM-provided paragraph structure tags.

        - Roles are aligned to role_spans in order.
        - Bullets are mapped within each role span by order of appearance.
        """
        bullet_mapping: Dict[str, int] = {}
        role_mapping: Dict[str, Tuple[int, int]] = {}

        try:
            role_spans = list(structure_tags.get("role_spans", []))
            role_spans = sorted(role_spans, key=lambda s: int(s.get("start_index", 0)))

            # Align roles to spans by order
            for role, span in zip(resume_ast.roles, role_spans):
                start_idx = int(span.get("start_index", 0))
                end_idx = int(span.get("end_index", start_idx))
                # Clamp to document bounds
                start_idx = max(0, min(start_idx, len(doc.paragraphs) - 1))
                end_idx = max(start_idx, min(end_idx, len(doc.paragraphs) - 1))
                role_mapping[role.id] = (start_idx, end_idx)

            # Build quick lookup for bullet-labeled paragraphs
            tag_list = structure_tags.get("paragraph_tags", [])
            bullet_indices_by_span: List[List[int]] = []
            for span in role_spans:
                s = int(span.get("start_index", 0))
                e = int(span.get("end_index", s))
                s = max(0, min(s, len(doc.paragraphs) - 1))
                e = max(s, min(e, len(doc.paragraphs) - 1))
                indices = [
                    int(t.get("paragraph_index"))
                    for t in tag_list
                    if t.get("label") == "bullet" and isinstance(t.get("paragraph_index"), int) and s <= int(t.get("paragraph_index")) <= e
                ]
                bullet_indices_by_span.append(sorted(indices))

            # Map ResumeAST bullets to bullet indices in order per role
            for role_idx, role in enumerate(resume_ast.roles):
                if role.id not in role_mapping:
                    continue
                indices = bullet_indices_by_span[role_idx] if role_idx < len(bullet_indices_by_span) else []
                role_bullets = [b for b in resume_ast.bullets if b.role_id == role.id]
                for b, p_idx in zip(role_bullets, indices):
                    bullet_mapping[b.id] = p_idx

            return bullet_mapping, role_mapping
        except Exception as e:
            self.logger.warning(f"Failed to build mappings from tags, falling back to heuristics: {e}")
            return self._create_bullet_mapping(doc, resume_ast), self._create_role_mapping(doc, resume_ast)
    
    def _create_bullet_mapping(self, doc: Document, resume_ast: ResumeAST) -> Dict[str, int]:
        """Create mapping between bullet IDs and paragraph indices."""
        bullet_mapping = {}
        
        # First pass: try exact matches
        if resume_ast.bullets:
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if self._is_bullet_paragraph(paragraph, text):
                    # Try exact match first
                    clean_text = self._normalize_bullet_text(text)
                    for bullet in resume_ast.bullets:
                        if self._normalize_bullet_text(bullet.text) == clean_text:
                            bullet_mapping[bullet.id] = i
                            self.logger.debug(f"Exact match for {bullet.id} at {i}: {text[:50]}...")
                            break
        
        # Second pass: try fuzzy matches for unmapped bullets
        if resume_ast.bullets:
            unmapped_bullets = [b for b in resume_ast.bullets if b.id not in bullet_mapping]
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if self._is_bullet_paragraph(paragraph, text):
                    # Skip if this paragraph is already mapped
                    if i in bullet_mapping.values():
                        continue
                    # Try fuzzy match
                    for bullet in unmapped_bullets:
                        if self._text_similarity(text, bullet.text) > 0.5:  # More lenient threshold
                            bullet_mapping[bullet.id] = i
                            self.logger.debug(f"Fuzzy match for {bullet.id} at {i}: {text[:50]}...")
                            break
        
        # If still missing mappings, create sequential IDs for unmapped bullets
        if not bullet_mapping or len(bullet_mapping) < len([p for p in doc.paragraphs if self._is_bullet_paragraph(p, p.text.strip())]):
            bullet_counter = max([int(bid.split('_')[1]) for bid in bullet_mapping.keys() if bid.startswith('bullet_') and bid.split('_')[1].isdigit()] + [0]) + 1
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if self._is_bullet_paragraph(paragraph, text) and i not in bullet_mapping.values():
                    bullet_id = f"bullet_{bullet_counter}"
                    bullet_mapping[bullet_id] = i
                    bullet_counter += 1
            
            self.logger.info(f"Added sequential mapping for {bullet_counter-1} unmapped bullets")
        
        return bullet_mapping
    
    def _is_bullet_paragraph(self, paragraph, text: str) -> bool:
        """Check if a paragraph is a bullet point using multiple detection methods."""
        # Method 1: Check for visible bullet characters
        if text.startswith('•') or text.startswith('-') or text.startswith('*') or text.startswith('◦'):
            return True
        
        # Method 2: Check Word's bullet formatting
        try:
            # Check if paragraph has bullet formatting
            if hasattr(paragraph, '_element') and paragraph._element is not None:
                # Look for bullet formatting in the paragraph element
                if paragraph._element.find('.//w:numPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
                    return True
                
                # Check for bullet list formatting
                if paragraph._element.find('.//w:listPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
                    return True
        except Exception:
            pass
        
        # Method 3: Check for common bullet patterns in text
        # Look for lines that start with common bullet indicators
        bullet_patterns = [
            r'^\s*[•\-*◦▪▫]\s+',  # Bullet characters
            r'^\s*\d+\.\s+',      # Numbered lists
            r'^\s*[a-z]\.\s+',    # Letter lists
            r'^\s*[A-Z]\.\s+',    # Capital letter lists
        ]
        
        import re
        for pattern in bullet_patterns:
            if re.match(pattern, text):
                return True
        
        # Method 4: Check paragraph style for bullet formatting
        try:
            if hasattr(paragraph, 'style') and paragraph.style:
                style_name = paragraph.style.name.lower()
                if any(bullet_indicator in style_name for bullet_indicator in ['list', 'bullet', 'number']):
                    return True
        except Exception:
            pass
        
        # Method 5: Check for indentation and short text (common bullet characteristics)
        # Only use this as a last resort and be more conservative
        if len(text) > 10 and len(text) < 200:  # Reasonable bullet length
            # Check if this looks like a bullet point based on content
            # But be more conservative - only if it's clearly a bullet
            if (text.startswith('•') or text.startswith('-') or text.startswith('*')):
                return True
            # Only check for action words if the text is very short (likely a bullet)
            if len(text) < 100 and any(keyword in text.lower() for keyword in ['developed', 'designed', 'implemented']):
                return True
        
        return False
    
    def _clean_bullet_text(self, text: str) -> str:
        """Clean up bullet text to ensure proper formatting."""
        # Remove any existing bullet characters at the start
        text = re.sub(r'^[•\-\*◦▪▫]\s*', '', text.strip())
        
        # Ensure text starts with a single bullet
        if not text.startswith('•'):
            text = f"• {text}"
        
        return text

    def _strip_leading_bullet(self, text: str) -> str:
        """Remove any visible bullet character/prefix from the beginning of text."""
        return re.sub(r'^[•\-\*◦▪▫]\s*', '', (text or '').strip())

    def _ensure_single_bullet_prefix(self, text: str) -> str:
        """Ensure the text has exactly one leading '• ' prefix (manual bullets style)."""
        return self._clean_bullet_text(text or '')
    
    # Removed _cleanup_duplicate_bullets - document editor should only modify what's explicitly requested
    
    # Removed _fix_missing_company_names - document editor should only modify what's explicitly requested
    
    # Removed company name helper methods - document editor should only modify what's explicitly requested
    
    def _ensure_bullet_formatting(self, paragraph, doc: Document, role_start: int, role_end: int):
        """Ensure the new paragraph has proper bullet formatting."""
        # Find existing bullet formatting in the role section
        for i in range(role_start, role_end + 1):
            existing_para = doc.paragraphs[i]
            if self._is_bullet_paragraph(existing_para, existing_para.text.strip()):
                # Copy the formatting from existing bullet
                try:
                    paragraph.style = existing_para.style
                    paragraph.alignment = existing_para.alignment
                    
                    # Copy font formatting if possible
                    if existing_para.runs and paragraph.runs:
                        existing_font = existing_para.runs[0].font
                        new_font = paragraph.runs[0].font
                        
                        # Copy font properties
                        new_font.name = existing_font.name
                        new_font.size = existing_font.size
                        new_font.bold = existing_font.bold
                        new_font.italic = existing_font.italic
                        
                        # Copy color if available
                        try:
                            if existing_font.color:
                                new_font.color.rgb = existing_font.color.rgb
                        except:
                            pass
                except Exception as e:
                    self.logger.warning(f"Could not copy all formatting: {e}")
                
                break
    
    def _create_role_mapping(self, doc: Document, resume_ast: ResumeAST) -> Dict[str, Tuple[int, int]]:
        """Create mapping between role IDs and document section boundaries."""
        role_mapping = {}
        
        # If ResumeAST has roles, try to map them using improved heuristics
        if resume_ast.roles:
            for role in resume_ast.roles:
                # Find role section boundaries
                start_idx, end_idx = self._find_role_section_improved(doc, role)
                if start_idx is not None and end_idx is not None:
                    role_mapping[role.id] = (start_idx, end_idx)
                    self.logger.debug(f"Mapped {role.id} to paragraphs {start_idx}-{end_idx}")
        
        # If no roles mapped from ResumeAST, create sequential mapping
        if not role_mapping:
            role_counter = 1
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                # Look for role-like patterns using improved detection
                if self._is_role_paragraph(text):
                    role_id = f"role_{role_counter}"
                    start_idx = i
                    end_idx = self._find_section_end(doc, i)
                    role_mapping[role_id] = (start_idx, end_idx)
                    role_counter += 1
            
            self.logger.info(f"Created sequential role mapping for {len(role_mapping)} roles")
        else:
            self.logger.info(f"Created role mapping for {len(role_mapping)} roles from ResumeAST")
        
        return role_mapping
    
    def _find_role_section_improved(self, doc: Document, role) -> Tuple[Optional[int], Optional[int]]:
        """Find the paragraph range for a specific role with improved heuristics."""
        
        # Strategy 1: Look for exact title match first
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if role.title.lower() in text.lower():
                # Found role title, now find section boundaries
                start_idx = i
                end_idx = self._find_section_end(doc, i)
                self.logger.debug(f"Found role '{role.title}' at paragraph {i}: {text[:50]}...")
                return start_idx, end_idx
        
        # Strategy 2: Look for company name match
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if role.company.lower() in text.lower():
                # Found company, look for the title above it (usually 1-2 paragraphs above)
                for j in range(max(0, i-2), i+1):
                    title_text = doc.paragraphs[j].text.strip()
                    if role.title.lower() in title_text.lower():
                        start_idx = j
                        end_idx = self._find_section_end(doc, j)
                        self.logger.debug(f"Found role '{role.title}' above company '{role.company}' at paragraph {j}")
                        return start_idx, end_idx
        
        # Strategy 3: Look for partial matches with better logic
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if len(text) < 10:  # Skip very short paragraphs
                continue
                
            # Check if this paragraph contains parts of both title and company
            title_words = [w.lower() for w in role.title.lower().split() if len(w) > 3]
            company_words = [w.lower() for w in role.company.lower().split() if len(w) > 3]
            
            title_match = any(word in text.lower() for word in title_words)
            company_match = any(word in text.lower() for word in company_words)
            
            if title_match and company_match:
                start_idx = i
                end_idx = self._find_section_end(doc, i)
                self.logger.debug(f"Found role by partial match at paragraph {i}: {text[:50]}...")
                return start_idx, end_idx
        
        # Strategy 4: Look for role patterns in sequence
        role_patterns = [
            'software engineer', 'developer', 'architect', 'manager', 'lead',
            'engineer', 'analyst', 'consultant', 'specialist'
        ]
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if any(pattern in text.lower() for pattern in role_patterns):
                # Check if this might be our role by looking at nearby text
                nearby_text = ' '.join([
                    doc.paragraphs[j].text.strip() 
                    for j in range(max(0, i-1), min(len(doc.paragraphs), i+2))
                ]).lower()
                
                if role.title.lower() in nearby_text or role.company.lower() in nearby_text:
                    start_idx = i
                    end_idx = self._find_section_end(doc, i)
                    self.logger.debug(f"Found role by pattern match at paragraph {i}: {text[:50]}...")
                    return start_idx, end_idx
        
        self.logger.warning(f"Could not find role section for {role.title} at {role.company}")
        return None, None
    
    def _find_section_end(self, doc: Document, start_idx: int) -> int:
        """Find where a role section ends - improved to handle section transitions."""
        # Major section headers that should terminate role sections
        section_headers = [
            'EDUCATION', 'INTERNSHIPS', 'SKILLS', 'PROJECTS', 'CERTIFICATIONS',
            'AWARDS', 'PUBLICATIONS', 'LANGUAGES', 'INTERESTS'
        ]
        
        # Look for either a new role or a major section
        for i in range(start_idx + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip()
            
            # Skip empty lines
            if not text:
                continue
                
            # Check for major section headers first
            if text.upper() in section_headers or any(text.upper().startswith(h) for h in section_headers):
                return i - 1
            
            # Then check for new role patterns
            if len(text) > 10:  # Skip very short lines
                # Role title keywords
                job_titles = [
                    'software engineer', 'developer', 'architect', 'manager', 'lead',
                    'engineer', 'analyst', 'consultant', 'specialist', 'intern'
                ]
                
                # Date patterns that often appear with role headers
                date_patterns = [
                    r'\d{2}/\d{2}\s*-\s*(present|current|\d{2}/\d{2})',  # MM/YY - Present
                    r'\d{4}\s*-\s*(present|current|\d{4})',              # YYYY - Present
                    r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+\d{4}',  # Month YYYY
                ]
                
                # Check for job title + date combination
                has_title = any(title in text.lower() for title in job_titles)
                has_date = any(re.search(pattern, text.lower()) for pattern in date_patterns)
                
                if has_title or (has_date and i > start_idx + 2):  # Allow some space for bullets
                    return i - 1
                
                # Also check next line for company name pattern
                if i + 1 < len(doc.paragraphs):
                    next_text = doc.paragraphs[i + 1].text.strip().lower()
                    if any(word in next_text for word in ['inc', 'corp', 'ltd', 'llc', 'systems', 'technologies']):
                        return i - 1
        
        return len(doc.paragraphs) - 1
    
    def _text_similarity(self, text1: str, text2: str) -> float:
        """Calculate text similarity between two strings."""
        # Simple similarity based on common words
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union)
    
    def _apply_single_edit(self, doc: Document, edit, bullet_mapping: Dict, 
                          role_mapping: Dict, resume_ast: ResumeAST):
        """Apply a single edit to the document."""
        
        # Handle both Plan object and dictionary
        if isinstance(edit, dict):
            edit_type = edit.get('type')
        else:
            edit_type = edit.type
        
        if edit_type == "modify_bullet":
            self._apply_modify_bullet(doc, edit, bullet_mapping)
        elif edit_type == "insert_bullet":
            self._apply_insert_bullet(doc, edit, bullet_mapping, role_mapping)
        elif edit_type == "upsert_skill":
            self._apply_upsert_skill(doc, edit, resume_ast)
        else:
            self.logger.warning(f"Unknown edit type: {edit_type}")
    
    def _apply_modify_bullet(self, doc: Document, edit, bullet_mapping: Dict):
        """Apply bullet modification edit."""
        # Handle both Plan object and dictionary
        if isinstance(edit, dict):
            bullet_id = edit.get('bullet_id')
            new_text = edit.get('new_text')
            resolved_index = edit.get('resolved_paragraph_index')
            pre_image_hash = edit.get('pre_image_hash')
        else:
            bullet_id = edit.bullet_id
            new_text = edit.new_text
            resolved_index = getattr(edit, 'resolved_paragraph_index', None)
            pre_image_hash = getattr(edit, 'pre_image_hash', None)
        
        if resolved_index is not None:
            paragraph_idx = int(resolved_index)
        else:
            if bullet_id not in bullet_mapping:
                self.logger.warning(f"Bullet {bullet_id} not found in mapping")
                return
            paragraph_idx = bullet_mapping[bullet_id]
        paragraph = doc.paragraphs[paragraph_idx]

        # Pre-image guard: ensure text hash matches if provided
        if pre_image_hash:
            import hashlib
            current_hash = hashlib.sha256((paragraph.text or '').encode('utf-8')).hexdigest()
            if current_hash != pre_image_hash:
                self.logger.warning(f"Pre-image hash mismatch at paragraph {paragraph_idx}; skipping modify")
                return
        
        # Preserve formatting while updating text
        original_text = paragraph.text
        
        # Preserve the original formatting
        original_style = paragraph.style
        original_alignment = paragraph.alignment
        original_font = None
        try:
            if paragraph.runs:
                original_font = paragraph.runs[0].font
        except:
            pass
        
        # Update the text with awareness of Word auto-bullets
        try:
            if self._has_auto_bullets(paragraph):
                # Strip any manual bullet prefix to avoid "• •" duplication
                paragraph.text = self._strip_leading_bullet(new_text or "")
            else:
                paragraph.text = new_text or ""
        except Exception:
            paragraph.text = new_text or ""
        
        # Restore formatting
        paragraph.style = original_style
        paragraph.alignment = original_alignment
        
        # Try to restore font formatting if possible
        if original_font and paragraph.runs:
            try:
                paragraph.runs[0].font.name = original_font.name
                paragraph.runs[0].font.size = original_font.size
                paragraph.runs[0].font.bold = original_font.bold
                paragraph.runs[0].font.italic = original_font.italic
            except:
                pass
        
        self.logger.info(f"Modified bullet {bullet_id}: '{original_text[:50]}...' → '{new_text[:50]}...'")
    
    def _apply_insert_bullet(self, doc: Document, edit, bullet_mapping: Dict, role_mapping: Dict):
        """Apply bullet insertion edit."""
        # Handle both Plan object and dictionary
        if isinstance(edit, dict):
            role_id = edit.get('role_id')
            text = edit.get('text')
            after_bullet_id = edit.get('after_bullet_id')
            resolved_after_index = edit.get('resolved_after_paragraph_index')
        else:
            role_id = edit.role_id
            text = edit.text
            after_bullet_id = edit.after_bullet_id
            resolved_after_index = getattr(edit, 'resolved_after_paragraph_index', None)
        
        if role_id not in role_mapping:
            self.logger.warning(f"Role {role_id} not found in mapping")
            return
        
        role_start, role_end = role_mapping[role_id]
        
        # Determine insertion position safely within the role span
        if resolved_after_index is not None:
            insert_idx = int(resolved_after_index) + 1
        elif after_bullet_id and after_bullet_id in bullet_mapping:
            insert_idx = bullet_mapping[after_bullet_id] + 1
        else:
            # Find the last bullet inside the role span
            last_bullet_idx = None
            for i in range(role_start, min(role_end + 1, len(doc.paragraphs))):
                p_txt = (doc.paragraphs[i].text or '').strip()
                if self._is_bullet_paragraph(doc.paragraphs[i], p_txt):
                    last_bullet_idx = i
            if last_bullet_idx is not None:
                insert_idx = last_bullet_idx + 1
            else:
                # No bullets in role yet: insert right after the role header/company line area
                insert_idx = min(role_start + 2, role_end + 1, len(doc.paragraphs))
        
        # De-duplicate within the role span
        if self._is_duplicate_bullet_in_span(doc, role_start, role_end, text or ""):
            self.logger.info("Skipping insert: duplicate bullet detected in role span")
            return

        # Insert new paragraph with proper bullet formatting
        new_paragraph = doc.add_paragraph()

        # Determine if this role section uses Word's automatic bullets
        uses_auto_bullets = self._role_uses_auto_bullets(doc, role_start, role_end)

        # Clean text appropriately based on bullet style
        if uses_auto_bullets:
            # Strip any leading manual bullet; Word will render bullet via paragraph formatting
            new_paragraph.text = self._strip_leading_bullet(text)
        else:
            # This document uses manual bullets in text; ensure a single leading bullet
            new_paragraph.text = self._ensure_single_bullet_prefix(text)
        
        # Move paragraph to correct position
        doc._body._body.insert(insert_idx, new_paragraph._element)
        
        # Apply consistent formatting from existing bullets in the role
        self._apply_bullet_formatting(new_paragraph, doc, role_start, role_end)

        # Ensure the new paragraph has bullet formatting (style + numbering)
        self._ensure_bullet_formatting(new_paragraph, doc, role_start, role_end)
        self._ensure_bullet_numbering(new_paragraph, doc, role_start, role_end)
        
        # Role-scoped dedupe after insertion to prevent duplicates
        try:
            self._dedupe_bullets_in_span(doc, role_start, role_end)
        except Exception:
            pass
        # No automatic cleanup - only apply the specific edit requested
        
        self.logger.info(f"Inserted new bullet in role {role_id}: '{text[:50]}...'")
    
    def _ensure_company_lines_preserved(self, doc: Document, resume_ast: ResumeAST, role_mapping: Dict[str, Tuple[int, int]], original_company_lines: List[str]):
        """Ensure all original company lines are present in the edited document.
        If missing, insert right after the role header start index.
        """
        self.logger.info("Ensuring company lines are preserved...")
        
        # Build a fast lookup of current document text lowercased
        doc_texts_lower = [(p.text or '').strip().lower() for p in doc.paragraphs]
        
        # Process each role in order
        for idx, role in enumerate(resume_ast.roles):
            if idx >= len(original_company_lines):
                continue
                
            company_line = (original_company_lines[idx] or '').strip()
            if not company_line:
                continue
                
            # Check if company line is already present
            company_lower = company_line.lower()
            if any(company_lower in text for text in doc_texts_lower):
                self.logger.debug(f"Company line already present for role {role.id}: {company_line}")
                continue
                
            # Company line is missing - need to insert it
            if role.id in role_mapping:
                start_idx, _ = role_mapping[role.id]
                self.logger.info(f"Inserting missing company line for role {role.id}: {company_line}")
                
                # Insert after the role header (start_idx)
                insert_idx = min(start_idx + 1, len(doc.paragraphs))
                
                # Create new paragraph with company line
                new_para = doc.add_paragraph()
                new_para.text = company_line
                
                # Copy formatting from nearby paragraphs if possible
                try:
                    if start_idx < len(doc.paragraphs):
                        source_para = doc.paragraphs[start_idx]
                        new_para.style = source_para.style
                        new_para.alignment = source_para.alignment
                        
                        # Copy font formatting if possible
                        if source_para.runs and new_para.runs:
                            source_font = source_para.runs[0].font
                            new_font = new_para.runs[0].font
                            
                            new_font.name = source_font.name
                            new_font.size = source_font.size
                            new_font.bold = source_font.bold
                            new_font.italic = source_font.italic
                            
                            try:
                                if source_font.color:
                                    new_font.color.rgb = source_font.color.rgb
                            except:
                                pass
                except Exception as e:
                    self.logger.warning(f"Could not copy formatting for company line: {e}")
                
                # Insert the paragraph at the correct position
                doc._body._body.insert(insert_idx, new_para._element)
                
                # Update the role mapping to account for the new paragraph
                for role_id, (r_start, r_end) in role_mapping.items():
                    if r_start >= insert_idx:
                        role_mapping[role_id] = (r_start + 1, r_end + 1)
                    elif r_end >= insert_idx:
                        role_mapping[role_id] = (r_start, r_end + 1)
                
                self.logger.info(f"Successfully inserted company line: {company_line}")
            else:
                self.logger.warning(f"No role mapping found for role {role.id}, cannot insert company line")
        
        self.logger.info("Company line preservation check completed")

    def _has_auto_bullets(self, paragraph) -> bool:
        """Return True if paragraph uses Word numbering/bullets (numPr present)."""
        try:
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            p_el = getattr(paragraph, '_element', None)
            return bool(p_el is not None and p_el.find('.//w:numPr', ns) is not None)
        except Exception:
            return False

    def _is_duplicate_bullet_in_span(self, doc: Document, role_start: int, role_end: int, candidate_text: str) -> bool:
        """Check for duplicate bullet (normalized) within a role span."""
        norm_new = self._normalize_bullet_text(candidate_text)
        if not norm_new:
            return False
        for i in range(role_start, min(role_end + 1, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            txt = (p.text or "").strip()
            if not txt:
                continue
            if self._is_bullet_paragraph(p, txt):
                if self._normalize_bullet_text(txt) == norm_new:
                    return True
        return False

    def _normalize_bullet_text(self, text: str) -> str:
        """Lowercased, bullet-prefix-stripped, single-spaced text for duplicate checks."""
        base = self._strip_leading_bullet(text or "").lower()
        import re as _re
        base = _re.sub(r"\s+", " ", base).strip()
        return base

    def _role_uses_auto_bullets(self, doc: Document, role_start: int, role_end: int) -> bool:
        """Detect if the role section uses Word's automatic bullet/numbering formatting."""
        try:
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            for i in range(role_start, min(role_end + 1, len(doc.paragraphs))):
                paragraph = doc.paragraphs[i]
                if not paragraph.text.strip():
                    continue
                # If it's a bullet paragraph and has numbering properties, it's auto bullets
                if self._is_bullet_paragraph(paragraph, paragraph.text.strip()):
                    p_el = getattr(paragraph, '_element', None)
                    if p_el is not None and p_el.find('.//w:numPr', namespaces) is not None:
                        return True
            return False
        except Exception:
            # If detection fails, default to False to avoid adding double bullets
            return False
    
    def _apply_upsert_skill(self, doc: Document, edit, resume_ast: ResumeAST):
        """Apply skill upsert edit."""
        # Handle both Plan object and dictionary
        if isinstance(edit, dict):
            value = edit.get('value')
            bucket = edit.get('bucket')
        else:
            value = edit.value
            bucket = edit.bucket
        
        # Find or create skills section
        skills_section = self._find_skills_section(doc)
        
        if skills_section:
            # Add skill to existing section
            self._add_skill_to_section(doc, skills_section, value, bucket)
        else:
            # Create new skills section
            self._create_skills_section(doc, bucket, value)
        
        self.logger.info(f"Upserted skill '{value}' in category '{bucket}'")
    
    def _find_skills_section(self, doc: Document) -> Optional[int]:
        """Find the skills section in the document."""
        skills_keywords = ['skills', 'technical skills', 'competencies', 'expertise']
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.lower()
            if any(keyword in text for keyword in skills_keywords):
                return i
        
        return None
    
    def _add_skill_to_section(self, doc: Document, section_idx: int, skill: str, category: str):
        """Add a skill to an existing skills section."""
        # Find the category within the section
        for i in range(section_idx, len(doc.paragraphs)):
            paragraph = doc.paragraphs[i]
            if category.lower() in paragraph.text.lower():
                # Add skill to this category
                if skill not in paragraph.text:
                    paragraph.text += f", {skill}"
                break
    
    def _create_skills_section(self, doc: Document, category: str, skill: str):
        """Create a new skills section."""
        # Add section header
        header = doc.add_paragraph()
        header.text = f"{category}:"
        try:
            # Try to use Heading 2 style if available
            header.style = 'Heading 2'
        except KeyError:
            # If style not found, apply direct formatting
            header.runs[0].font.bold = True
            header.runs[0].font.size = 280000  # 14pt
        
        # Add skills
        skills_para = doc.add_paragraph()
        skills_para.text = skill
    
    def _apply_bullet_formatting(self, paragraph, doc: Document, role_start: int, role_end: int):
        """Apply consistent bullet formatting to a paragraph."""
        # Find existing bullet formatting in the role section
        for i in range(role_start, role_end + 1):
            existing_para = doc.paragraphs[i]
            if existing_para.text.strip().startswith('•'):
                # Copy formatting from existing bullet
                paragraph.style = existing_para.style
                paragraph.alignment = existing_para.alignment
                break

    def _ensure_bullet_numbering(self, paragraph, doc: Document, role_start: int, role_end: int):
        """Clone Word numbering (w:numPr) from a nearby bullet to the new paragraph if present."""
        try:
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            # Find a paragraph within the role that has numPr
            donor = None
            for i in range(role_start, role_end + 1):
                p = doc.paragraphs[i]
                p_el = getattr(p, '_element', None)
                if p_el is not None and p_el.find('.//w:numPr', ns) is not None:
                    donor = p
                    break
            if donor is None:
                return
            donor_el = getattr(donor, '_element', None)
            target_el = getattr(paragraph, '_element', None)
            if donor_el is None or target_el is None:
                return
            # Remove any existing numPr on target
            existing = target_el.find('.//w:numPr', ns)
            if existing is not None:
                parent = existing.getparent()
                if parent is not None:
                    parent.remove(existing)
            # Clone donor numPr into target pPr
            from copy import deepcopy
            donor_numpr = donor_el.find('.//w:numPr', ns)
            if donor_numpr is None:
                return
            # Ensure target has pPr
            ppr = target_el.find('.//w:pPr', ns)
            if ppr is None:
                # create pPr
                from lxml import etree
                ppr = etree.Element(f"{{{ns['w']}}}pPr")
                target_el.insert(0, ppr)
            ppr.append(deepcopy(donor_numpr))
        except Exception as e:
            self.logger.debug(f"Could not clone numbering: {e}")

    def _dedupe_bullets_in_span(self, doc: Document, role_start: int, role_end: int):
        """Remove later duplicate bullets within a role span (case/space-insensitive)."""
        seen = set()
        to_remove = []
        for i in range(role_start, min(role_end + 1, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            t = (p.text or '').strip()
            if not t:
                continue
            if self._is_bullet_paragraph(p, t):
                norm = self._normalize_bullet_text(t)
                if norm in seen:
                    to_remove.append(i)
                else:
                    seen.add(norm)
        # Remove from bottom up
        for i in reversed(to_remove):
            elem = doc.paragraphs[i]._element
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
    
    def _generate_output_path(self, original_path: str) -> str:
        """Generate output path for modified document in resumes directory."""
        # Ensure resumes directory exists
        resumes_dir = Path("resumes")
        resumes_dir.mkdir(exist_ok=True)
        
        # Get original filename and generate enhanced name
        path = Path(original_path)
        output_name = f"{path.stem}_enhanced{path.suffix}"
        output_path = resumes_dir / output_name
        
        # Ensure unique filename
        counter = 1
        while output_path.exists():
            output_name = f"{path.stem}_enhanced_{counter}{path.suffix}"
            output_path = resumes_dir / output_name
            counter += 1
        
        return str(output_path)
    
    def _is_role_paragraph(self, text: str) -> bool:
        """Check if a paragraph represents a job role - simplified to avoid false positives."""
        # Only check for clear job titles, not company/location/date patterns
        job_titles = [
            'software engineer', 'developer', 'architect', 'manager', 'lead', 'consultant',
            'analyst', 'specialist', 'coordinator', 'director', 'officer', 'intern',
            'tutor', 'associate', 'senior', 'principal', 'staff', 'chief'
        ]
        
        # Check if text contains a job title
        if any(title in text.lower() for title in job_titles):
            return True
        
        return False
