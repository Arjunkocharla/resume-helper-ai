#!/usr/bin/env python3
"""LLM Direct Editor Agent - Uses LLM to apply improvement plans directly to resume content."""

import logging
import os
from typing import Dict, List, Optional, Union

import anthropic
from docx import Document

from ..contracts import Plan, ResumeAST, validate_or_raise

logger = logging.getLogger(__name__)


class LLMDirectEditorAgent:
    """Agent responsible for using LLM to directly apply improvement plans to resume content."""
    
    def __init__(self, client: anthropic.Anthropic):
        self.client = client
        self.logger = logging.getLogger(__name__)
    
    def enhance_resume(self, resume_path: str, plan: Union[Plan, dict], resume_ast: ResumeAST, 
                      jd_context: str = "") -> str:
        """
        Enhance resume using LLM to apply improvement plan directly.
        
        Args:
            resume_path: Path to original resume file (PDF or DOCX)
            plan: Improvement plan with edits to apply
            resume_ast: Parsed resume structure for context
            jd_context: Job description context for better understanding
            
        Returns:
            str: Enhanced resume content as text
            
        Raises:
            ValueError: If enhancement cannot be completed
            FileNotFoundError: If resume file doesn't exist
        """
        try:
            if not os.path.exists(resume_path):
                raise FileNotFoundError(f"Resume file not found: {resume_path}")
            
            # Handle both Plan objects and dictionaries
            if hasattr(plan, 'edits'):
                edits = plan.edits
            else:
                edits = plan.get('edits', [])
            
            self.logger.info(f"Enhancing resume with {len(edits)} improvements")
            
            # Extract resume content
            resume_content = self._extract_resume_content(resume_path)
            
            # Apply improvements using LLM
            enhanced_content = self._apply_improvements_with_llm(
                resume_content, plan, resume_ast, jd_context
            )
            
            self.logger.info("Successfully enhanced resume using LLM")
            return enhanced_content
            
        except Exception as e:
            self.logger.error(f"Failed to enhance resume: {e}")
            raise
    
    def _extract_resume_content(self, resume_path: str) -> str:
        """Extract text content from resume file (PDF or DOCX)."""
        file_ext = os.path.splitext(resume_path)[1].lower()
        
        if file_ext == '.docx':
            return self._extract_docx_content(resume_path)
        elif file_ext == '.pdf':
            return self._extract_pdf_content(resume_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_docx_content(self, docx_path: str) -> str:
        """Extract text content from DOCX file."""
        doc = Document(docx_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        return '\n'.join(content)
    
    def _extract_pdf_content(self, pdf_path: str) -> str:
        """Extract text content from PDF file."""
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = []
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content.append(text.strip())
                
                return '\n\n'.join(content)
                
        except ImportError:
            raise ValueError("PyPDF2 required for PDF processing")
    
    def _apply_improvements_with_llm(self, resume_content: str, plan: Plan, 
                                   resume_ast: ResumeAST, jd_context: str) -> str:
        """Use LLM to apply all improvements to resume content."""
        
        # Create comprehensive prompt for LLM
        system_prompt = self._create_system_prompt()
        user_prompt = self._create_user_prompt(resume_content, plan, resume_ast, jd_context)
        
        try:
            response = self.client.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=4000,
                temperature=0.1,
                system=system_prompt,
                messages=[
                    {"role": "user", "content": user_prompt}
                ]
            )
            
            enhanced_content = response.content[0].text.strip()
            
            # Validate that we got enhanced content
            if len(enhanced_content) < len(resume_content) * 0.5:
                raise ValueError("LLM returned insufficient content")
            
            return enhanced_content
            
        except Exception as e:
            self.logger.error(f"LLM enhancement failed: {e}")
            raise
    
    def _create_system_prompt(self) -> str:
        """Create system prompt for LLM enhancement."""
        return """You are an expert resume editor and career strategist. Your task is to enhance resumes by applying specific improvement plans.

CRITICAL RULES:
1. ONLY modify what the improvement plan specifies
2. Preserve ALL original content, structure, and formatting
3. Keep the same tone, style, and professional language
4. Do NOT add information that wasn't in the original resume
5. Do NOT fabricate experiences, skills, or achievements
6. Only enhance existing content with better wording and relevant keywords
7. Maintain exact same sections, headers, and organization
8. Apply all edits from the improvement plan exactly as specified

Your job is to make the resume stronger for the target job while staying completely honest and factual."""

    def _create_user_prompt(self, resume_content: str, plan: Plan, 
                           resume_ast: ResumeAST, jd_context: str) -> str:
        """Create user prompt for LLM enhancement."""
        
        # Format the improvement plan for clarity
        plan_summary = self._format_plan_summary(plan)
        
        # Create context about the resume structure
        resume_context = self._format_resume_context(resume_ast)
        
        prompt = f"""Enhance this resume by applying the improvement plan:

IMPROVEMENT PLAN:
{plan_summary}

RESUME STRUCTURE CONTEXT:
{resume_context}

JOB CONTEXT:
{jd_context}

ORIGINAL RESUME CONTENT:
{resume_content}

INSTRUCTIONS:
1. Apply ALL edits from the improvement plan exactly as specified
2. For bullet modifications: Replace the specified bullet with the enhanced version
3. For bullet insertions: Add new bullets in the correct positions
4. For skill updates: Add specified skills to appropriate sections
5. Preserve ALL original content, structure, and formatting
6. Keep the same professional tone and style
7. Return the COMPLETE enhanced resume

Return the complete enhanced resume with all improvements applied:"""

        return prompt
    
    def _format_plan_summary(self, plan: Union[Plan, dict]) -> str:
        """Format improvement plan for LLM prompt."""
        summary = []
        
        # Handle both Plan objects and dictionaries
        if hasattr(plan, 'edits'):
            edits = plan.edits
        else:
            edits = plan.get('edits', [])
        
        for i, edit in enumerate(edits, 1):
            # Handle both object and dictionary edits
            if hasattr(edit, 'type'):
                edit_type = edit.type
                bullet_id = getattr(edit, 'bullet_id', None)
                new_text = getattr(edit, 'new_text', None)
                role_id = getattr(edit, 'role_id', None)
                text = getattr(edit, 'text', None)
                after_bullet_id = getattr(edit, 'after_bullet_id', None)
                value = getattr(edit, 'value', None)
                bucket = getattr(edit, 'bucket', None)
            else:
                edit_type = edit.get('type')
                bullet_id = edit.get('bullet_id')
                new_text = edit.get('new_text')
                role_id = edit.get('role_id')
                text = edit.get('text')
                after_bullet_id = edit.get('after_bullet_id')
                value = edit.get('value')
                bucket = edit.get('bucket')
            
            if edit_type == "modify_bullet":
                summary.append(f"{i}. MODIFY BULLET: {bullet_id}")
                summary.append(f"   New text: {new_text}")
            elif edit_type == "insert_bullet":
                summary.append(f"{i}. INSERT BULLET in role: {role_id}")
                summary.append(f"   Text: {text}")
                if after_bullet_id:
                    summary.append(f"   After: {after_bullet_id}")
            elif edit_type == "upsert_skill":
                summary.append(f"{i}. ADD SKILL: {value} to {bucket}")
            
            summary.append("")
        
        return "\n".join(summary)
    
    def _format_resume_context(self, resume_ast: ResumeAST) -> str:
        """Format resume structure context for LLM prompt."""
        context = []
        
        context.append("RESUME SECTIONS:")
        for section in resume_ast.sections:
            context.append(f"  - {section.name}")
        
        context.append("\nROLES:")
        for role in resume_ast.roles:
            context.append(f"  - {role.title} at {role.company} ({role.date_range})")
        
        context.append("\nBULLETS:")
        for bullet in resume_ast.bullets:
            context.append(f"  - {bullet.id}: {bullet.text[:80]}...")
        
        context.append("\nSKILLS CATEGORIES:")
        for category, skills in resume_ast.skills_buckets.items():
            context.append(f"  - {category}: {', '.join(skills[:3])}...")
        
        return "\n".join(context)
    
    def create_enhanced_docx(self, enhanced_content: str, output_path: str) -> str:
        """
        Convert enhanced text content back to DOCX format.
        
        Args:
            enhanced_content: Enhanced resume content as text
            output_path: Path for output DOCX file
            
        Returns:
            str: Path to created DOCX file
        """
        try:
            doc = Document()
            
            # Parse enhanced content and recreate document structure
            lines = enhanced_content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Detect headers (all caps, short lines)
                if line.isupper() and len(line) < 50:
                    doc.add_heading(line, level=1)
                # Detect sub-headers (contains common section keywords)
                elif any(keyword in line.lower() for keyword in ['experience', 'skills', 'education', 'projects', 'certifications']):
                    doc.add_heading(line, level=2)
                # Regular content
                else:
                    doc.add_paragraph(line)
            
            # Save the document
            doc.save(output_path)
            self.logger.info(f"Enhanced DOCX saved to: {output_path}")
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to create enhanced DOCX: {e}")
            raise
