#!/usr/bin/env python3
"""LLMDocumentEditorAgent - Uses Claude 3.5 Sonnet for intelligent document editing.

This agent extracts document content, sends it to Claude for editing instructions,
then applies the changes programmatically while preserving formatting.
"""

import logging
import os
import re
from typing import Dict, Optional, List
from pathlib import Path
import anthropic

from docx import Document
from ..contracts import Plan, ResumeAST

logger = logging.getLogger(__name__)


class LLMDocumentEditorAgent:
    """Agent that uses Claude 3.5 Sonnet for intelligent document editing."""
    
    def __init__(self, anthropic_client):
        """Initialize with Anthropic client."""
        self.client = anthropic_client
        self.logger = logging.getLogger(__name__)
    
    def apply_plan(self, docx_path: str, plan: Plan, resume_ast: ResumeAST, structure_tags: Optional[Dict] = None) -> str:
        """
        Apply improvement plan using Claude 3.5 Sonnet for direct document editing.
        
        Args:
            docx_path: Path to original DOCX file
            plan: Improvement plan with edits to apply
            resume_ast: Parsed resume structure for context
            structure_tags: Optional structure tags (not used in direct LLM editing)
            
        Returns:
            str: Path to modified DOCX file
        """
        try:
            if not os.path.exists(docx_path):
                raise FileNotFoundError(f"DOCX file not found: {docx_path}")
            
            self.logger.info(f"Applying plan with direct LLM editing to {docx_path}")
            
            # Extract document content
            doc_content = self._extract_document_content(docx_path)
            
            # Send document + plan to LLM for direct editing
            edited_content = self._edit_document_with_llm(doc_content, plan, resume_ast)
            
            # Save the edited content as a new DOCX
            output_path = self._generate_output_path(docx_path)
            self._save_edited_content_to_docx(edited_content, output_path)
            
            self.logger.info(f"Successfully applied plan using direct LLM editing. Modified file saved to: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to apply plan with direct LLM editing: {e}")
            raise
    
    def _extract_document_content(self, docx_path: str) -> str:
        """Extract text content from document while preserving structure."""
        doc = Document(docx_path)
        content_lines = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                # Add paragraph index for reference
                content_lines.append(f"[{i:2}] {text}")
            else:
                content_lines.append(f"[{i:2}] <empty>")
        
        return "\n".join(content_lines)
    
    def _edit_document_with_llm(self, doc_content: str, plan: Plan, resume_ast: ResumeAST) -> str:
        """Send document content + plan to Claude 3.5 Sonnet for direct editing."""
        
        # Create edits summary
        edits_summary = []
        for edit in plan.edits:
            if hasattr(edit, 'type'):
                edit_type = edit.type
                if edit_type == 'modify_bullet':
                    edits_summary.append(f"MODIFY: Change bullet to: {edit.new_text}")
                elif edit_type == 'insert_bullet':
                    edits_summary.append(f"INSERT: Add bullet: {edit.text}")
                elif edit_type == 'upsert_skill':
                    edits_summary.append(f"ADD SKILL: {edit.value} to {edit.bucket}")
            else:
                edit_type = edit.get('type')
                if edit_type == 'modify_bullet':
                    edits_summary.append(f"MODIFY: Change bullet to: {edit.get('new_text')}")
                elif edit_type == 'insert_bullet':
                    edits_summary.append(f"INSERT: Add bullet: {edit.get('text')}")
                elif edit_type == 'upsert_skill':
                    edits_summary.append(f"ADD SKILL: {edit.get('value')} to {edit.get('bucket')}")
        
        prompt = f"""
You are an expert document editor. You will receive a resume document and a list of specific edits to make.

DOCUMENT CONTENT:
{doc_content}

REQUESTED CHANGES:
{chr(10).join(edits_summary)}

RESUME CONTEXT:
- Roles: {len(resume_ast.roles)} job positions
- Bullets: {len(resume_ast.bullets)} bullet points
- Skills: {len(plan.edits)} improvements requested

TASK:
Apply the requested changes directly to the document content above. Return the COMPLETE edited document.

CRITICAL REQUIREMENTS:
- PRESERVE ALL FORMATTING (fonts, sizes, colors, alignment, bullet styles)
- MAINTAIN DOCUMENT STRUCTURE (don't move sections or change layout)
- ONLY MODIFY REQUESTED CONTENT
- PRESERVE COMPANY LINES and role headers
- MAINTAIN bullet formatting and numbering
- Keep paragraph indices [XX] for reference

Return ONLY the edited document content with the same format as input, including paragraph indices.
"""
        
        try:
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4096,
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            )
            
            # Return the edited content directly
            edited_content = response.content[0].text
            self.logger.info(f"Received edited document content from LLM")
            return edited_content
            
        except Exception as e:
            self.logger.error(f"Failed to get edited document from LLM: {e}")
            raise Exception(f"LLM document editing failed: {e}")
    
    def _save_edited_content_to_docx(self, edited_content: str, output_path: str):
        """Save the LLM-edited content back to a DOCX file."""
        try:
            # Create a new document
            doc = Document()
            
            # Parse the edited content line by line
            lines = edited_content.strip().split('\n')
            
            for line in lines:
                # Skip empty lines
                if not line.strip():
                    continue
                
                # Extract text content (remove paragraph index if present)
                text_match = re.match(r'^\[\d+\]\s*(.*)$', line)
                if text_match:
                    text = text_match.group(1).strip()
                else:
                    text = line.strip()
                
                # Skip empty text after extraction
                if not text or text == '<empty>':
                    continue
                
                # Add paragraph to document
                if text:
                    doc.add_paragraph(text)
            
            # Save the document
            doc.save(output_path)
            self.logger.info(f"Saved edited content to DOCX: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to save edited content to DOCX: {e}")
            raise Exception(f"Failed to save edited content: {e}")
    
    def _generate_output_path(self, original_path: str) -> str:
        """Generate output path for modified document."""
        path = Path(original_path)
        output_name = f"{path.stem}_llm_enhanced{path.suffix}"
        output_path = path.parent / output_name
        
        # Ensure unique filename
        counter = 1
        while output_path.exists():
            output_name = f"{path.stem}_llm_enhanced_{counter}{path.suffix}"
            output_path = path.parent / output_name
            counter += 1
        
        return str(output_path)
