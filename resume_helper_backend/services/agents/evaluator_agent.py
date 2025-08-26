#!/usr/bin/env python3
"""EvaluatorAgent - deterministic checks after editing to ensure quality.

Checks include:
- Duplicate bullets within role spans
- Double bullet characters (e.g., "• •")
- Bullet formatting (missing numbering for bullets)
- Company lines preserved (count and text deltas)
Produces a structured report for the orchestrator.
"""

from __future__ import annotations

import logging
from typing import Dict, List, Tuple
from docx import Document
import re

logger = logging.getLogger(__name__)


class EvaluatorAgent:
    def evaluate(self,
                 original_docx_path: str,
                 enhanced_docx_path: str,
                 original_tags: Dict,
                 enhanced_tags: Dict) -> Dict:
        try:
            orig_doc = Document(original_docx_path)
            enh_doc = Document(enhanced_docx_path)

            report: Dict = {
                "passed": True,
                "issues": [],
                "duplicates": [],
                "double_bullets": [],
                "lost_bullets": [],
                "company_changes": [],
                "orphaned_bullets": [],
                "structural_issues": [],
            }

            # Duplicate bullets: check for identical bullet content within role spans
            duplicate_bullets = []
            role_mapping = {}
            for t in enhanced_tags.get("paragraph_tags", []):
                if t.get("label") == "role_header":
                    role_id = t.get("role_id")
                    if role_id:
                        start_idx = int(t.get("paragraph_index"))
                        # Find the end of the role span
                        end_idx = start_idx
                        for j in range(start_idx + 1, len(enh_doc.paragraphs)):
                            if enh_doc.paragraphs[j].text.strip().startswith("•") or enh_doc.paragraphs[j].text.strip().startswith("-") or enh_doc.paragraphs[j].text.strip().startswith("*"):
                                end_idx = j
                                break
                        role_mapping[role_id] = (start_idx, end_idx)

            for role_id, (start_idx, end_idx) in role_mapping.items():
                role_bullets = []
                for i in range(start_idx, end_idx + 1):
                    if i < len(enh_doc.paragraphs):
                        para = enh_doc.paragraphs[i]
                        if self._is_bullet_paragraph(para, para.text.strip()):
                            bullet_text = para.text.strip()
                            if bullet_text:
                                role_bullets.append((i, bullet_text))
                
                # Check for duplicates within this role
                seen_bullets = set()
                for idx, bullet_text in role_bullets:
                    # Normalize text for comparison (remove extra whitespace, bullet chars)
                    normalized = self._normalize_bullet_text(bullet_text)
                    if normalized in seen_bullets:
                        duplicate_bullets.append({
                            "role_id": role_id,
                            "paragraph_index": idx,
                            "text": bullet_text[:100] + "..." if len(bullet_text) > 100 else bullet_text,
                            "duplicate_of": next(i for i, t in role_bullets if self._normalize_bullet_text(t) == normalized)
                        })
                    else:
                        seen_bullets.add(normalized)
            
            if duplicate_bullets:
                report["duplicate_bullets"] = duplicate_bullets

            # Double bullet characters anywhere
            dbl_pat = re.compile(r"^\s*[•\-*]\s*[•\-*]\s+")
            for i, p in enumerate(enh_doc.paragraphs):
                if dbl_pat.match(p.text or ""):
                    report["double_bullets"].append({"paragraph_index": i, "text": p.text.strip()})

            # Lost bullets: paragraphs tagged as bullet but lacking numbering & visible bullet symbol
            enh_tag_map = {int(t.get("paragraph_index")): t.get("label") for t in enhanced_tags.get("paragraph_tags", []) if isinstance(t.get("paragraph_index"), int)}
            for i, p in enumerate(enh_doc.paragraphs):
                if enh_tag_map.get(i) == "bullet":
                    txt = (p.text or "").lstrip()
                    if not self._has_numpr(p) and not txt.startswith("•") and not txt.startswith("-") and not txt.startswith("*"):
                        report["lost_bullets"].append({"paragraph_index": i, "text": p.text.strip()})

            # Company line changes: ensure each original company line is present (ignore duplicates/order)
            orig_companies = self._collect_label_texts(original_tags, orig_doc, label="company_line")
            enh_companies = self._collect_label_texts(enhanced_tags, enh_doc, label="company_line")
            missing_companies = []
            enh_set_norm = { (c or '').strip().lower() for c in enh_companies }
            for c in orig_companies:
                c_norm = (c or '').strip().lower()
                if not c_norm:
                    continue
                # Present by tag set or by raw text search
                in_tags = c_norm in enh_set_norm
                in_text = any(((p.text or '').strip().lower().find(c_norm) >= 0) for p in enh_doc.paragraphs)
                if not (in_tags or in_text):
                    missing_companies.append(c)
            if missing_companies:
                report["company_changes"].append({
                    "missing": missing_companies,
                    "original_all": orig_companies,
                    "enhanced_tagged": enh_companies
                })

            # Drift check: non-bullet paragraphs should not change (if mapping present in tags)
            try:
                # Accept optional paragraph hashes from orchestrator via tags
                orig_hashes = original_tags.get('paragraph_hashes')
                if isinstance(orig_hashes, list) and len(orig_hashes) == len(orig_doc.paragraphs):
                    import hashlib
                    for i, p in enumerate(enh_doc.paragraphs):
                        lbl = None
                        for t in enhanced_tags.get('paragraph_tags', []):
                            if t.get('paragraph_index') == i:
                                lbl = t.get('label')
                                break
                        if lbl not in ("bullet", "skills_line"):
                            cur_hash = hashlib.sha256((p.text or '').encode('utf-8')).hexdigest()
                            if cur_hash != orig_hashes[i]:
                                report.setdefault('drift', []).append({
                                    'paragraph_index': i,
                                    'original_hash': orig_hashes[i],
                                    'current_hash': cur_hash
                                })
            except Exception:
                pass

            # Check for orphaned bullets (bullets not preceded by role headers within a stricter window)
            for i, p in enumerate(enh_doc.paragraphs):
                if enh_tag_map.get(i) == "bullet" or self._looks_like_bullet(enh_doc, i, p.text.strip()):
                    # Look backwards for a role header within 20 paragraphs, with heuristics
                    found_role = False
                    look_start = max(0, i-20)
                    job_title_keywords = [
                        'engineer', 'developer', 'architect', 'manager', 'lead', 'consultant',
                        'analyst', 'specialist', 'coordinator', 'director', 'officer', 'intern',
                        'tutor', 'associate', 'senior', 'principal', 'staff', 'chief'
                    ]
                    for j in range(look_start, i):
                        lbl = enh_tag_map.get(j)
                        if lbl in ["role_header", "company_line"]:
                            found_role = True
                            break
                        prev_text = (enh_doc.paragraphs[j].text or '').strip()
                        # Heuristic: treat as role header if contains common job titles
                        if any(k in prev_text.lower() for k in job_title_keywords):
                            found_role = True
                            break
                        # Heuristic: treat as company line if contains comma-separated location hints
                        if re.search(r",\s*[A-Z][a-zA-Z]+(,\s*[A-Z]{2})?$", prev_text):
                            found_role = True
                            break
                    if not found_role:
                        report["orphaned_bullets"].append({"paragraph_index": i, "text": p.text.strip()})

            # Check for structural issues (section headers in wrong places, missing role info)
            for i, p in enumerate(enh_doc.paragraphs):
                text = p.text.strip()
                if text in ["INTERNSHIPS", "WORK EXPERIENCE", "EDUCATION"] and i > 0:
                    # Check if this section header appears after bullets (indicating misplacement)
                    prev_text = enh_doc.paragraphs[i-1].text.strip()
                    if prev_text.startswith("•"):
                        report["structural_issues"].append({
                            "paragraph_index": i, 
                            "text": text,
                            "issue": "section_header_after_bullet"
                        })

            # Overall pass
            if (report["duplicates"] or report["double_bullets"] or report["lost_bullets"] or 
                report["company_changes"] or report["orphaned_bullets"] or report["structural_issues"] or report.get('drift')):
                report["passed"] = False
                if report["duplicates"]:
                    report["issues"].append("duplicate_bullets")
                if report["double_bullets"]:
                    report["issues"].append("double_bullet_chars")
                if report["lost_bullets"]:
                    report["issues"].append("lost_bullet_formatting")
                if report["company_changes"]:
                    report["issues"].append("company_lines_changed")
                if report["orphaned_bullets"]:
                    report["issues"].append("orphaned_bullets")
                if report["structural_issues"]:
                    report["issues"].append("structural_issues")
                if report.get('drift'):
                    report["issues"].append("non_allowlisted_drift")

            return report
        except Exception as e:
            logger.error(f"Evaluation failed: {e}")
            return {"passed": False, "issues": ["evaluator_error"], "error": str(e)}

    def _looks_like_bullet(self, doc: Document, idx: int, text: str) -> bool:
        if self._has_numpr(doc.paragraphs[idx]):
            return True
        return bool(text.startswith("•") or text.startswith("-") or text.startswith("*"))

    def _is_bullet_paragraph(self, paragraph, text: str) -> bool:
        """Determines if a paragraph is a bullet point based on its text content."""
        # Simple heuristic: check if the text starts with a bullet character (•, -, *)
        # and if it doesn't have a numbering property.
        # This is a simplification; a more robust check might involve parsing the paragraph's run properties.
        return bool(text.startswith("•") or text.startswith("-") or text.startswith("*")) and not self._has_numpr(paragraph)

    def _normalize_bullet_text(self, text: str) -> str:
        """Normalizes bullet text for duplicate detection."""
        # Remove bullet characters and extra whitespace
        base = re.sub(r'^[•\-*]\s*', '', (text or '').lower()).strip()
        base = re.sub(r'\s+', ' ', base)
        return base

    def _has_numpr(self, paragraph) -> bool:
        try:
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            p_el = getattr(paragraph, '_element', None)
            return bool(p_el is not None and p_el.find('.//w:numPr', ns) is not None)
        except Exception:
            return False

    def _collect_label_texts(self, tags: Dict, doc: Document, label: str) -> List[str]:
        results: List[str] = []
        for t in tags.get("paragraph_tags", []):
            if t.get("label") == label and isinstance(t.get("paragraph_index"), int):
                i = int(t.get("paragraph_index"))
                if 0 <= i < len(doc.paragraphs):
                    results.append((doc.paragraphs[i].text or '').strip())
        return results
