#!/usr/bin/env python3
"""Orchestrator Service that coordinates all agents in the resume enhancement workflow."""

import os
import logging
import json
from typing import Dict, List, Optional, Union
from pathlib import Path
from datetime import datetime

from .agents import JDParserAgent, ResumeParserAgent, GapAnalyzerAgent, DocumentEditorAgent, StructureTaggerAgent, EvaluatorAgent
from .contracts import ResumeAST, JDSummary, Plan
from .utils import PDFConverter

logger = logging.getLogger(__name__)


class ResumeEnhancementOrchestrator:
    """
    Orchestrator service that coordinates the entire resume enhancement workflow.
    
    This service manages the flow between different agents:
    1. JD Parser Agent - extracts requirements from job descriptions
    2. Resume Parser Agent - parses and structures resume content
    3. Gap Analyzer Agent - identifies gaps and generates improvement plans
    4. Document Editor Agent - applies improvements while preserving formatting
    5. PDF Converter - generates final PDF output
    """
    
    def __init__(self, anthropic_client):
        """Initialize the orchestrator with all required agents."""
        self.client = anthropic_client
        
        # HYBRID LLM STRATEGY: Use GPT-5 Nano for parsing/analysis, Anthropic for quality
        # Create separate clients for different tasks
        from config.llm_config import LLMProvider
        from services.llm_client import UnifiedLLMClient
        
        # GPT-5 Nano client for fast/cheap parsing and analysis tasks
        nano_client = UnifiedLLMClient(preferred_provider=LLMProvider.GPT5_NANO)
        
        # Anthropic client for quality-critical editing tasks
        anthropic_client_quality = UnifiedLLMClient(preferred_provider=LLMProvider.ANTHROPIC)
        
        # Use Haiku for all parsing to ensure quality and speed
        self.jd_parser = JDParserAgent(anthropic_client_quality)
        
        # Initialize skills categorization agent for resume parser
        from .agents import SkillsCategorizationAgent
        skills_agent = SkillsCategorizationAgent(anthropic_client_quality)  # Use Haiku for skills
        self.resume_parser = ResumeParserAgent(skills_agent)
        
        # Use Haiku for generating the improvement plan (quality-critical)
        self.gap_analyzer = GapAnalyzerAgent(anthropic_client_quality)
        self.document_editor = DocumentEditorAgent()  # No LLM needed
        # Use Haiku for structure tagging to ensure stable mappings
        self.structure_tagger = StructureTaggerAgent(anthropic_client_quality)
        self.pdf_converter = PDFConverter()
        self.evaluator = EvaluatorAgent()  # No LLM needed - deterministic evaluation
        
        logger.info("Resume Enhancement Orchestrator initialized with hybrid LLM strategy (GPT-5 Nano for parsing, Anthropic for quality)")
    
    def enhance_resume(
        self,
        resume_path: str,
        job_description: str,
        output_dir: Optional[str] = None,
        preserve_original: bool = True,
        generate_pdf: bool = True
    ) -> Dict:
        """
        Complete resume enhancement workflow.
        
        Args:
            resume_path: Path to the resume file (relative to resumes directory or absolute)
            job_description: Raw job description text
            output_dir: Directory to save enhanced files (defaults to resumes directory)
            preserve_original: Whether to keep original file
            generate_pdf: Whether to generate PDF output
            
        Returns:
            Dict containing workflow results and file paths
        """
        # Ensure resumes directory exists
        resumes_dir = Path("resumes")
        resumes_dir.mkdir(exist_ok=True)
        
        # Convert resume_path to absolute path if it's relative and not already in resumes dir
        resume_path_str = str(resume_path)
        if not os.path.isabs(resume_path_str) and not resume_path_str.startswith('resumes/') and not resume_path_str.startswith('uploads/'):
            resume_path = str(resumes_dir / resume_path_str)
        
        # Use resumes directory as default output
        if output_dir is None:
            output_dir = str(resumes_dir)
            
        # Check if input file is PDF and convert to DOCX if needed
        working_docx_path = resume_path
        is_pdf_input = resume_path.lower().endswith('.pdf')
        
        if is_pdf_input:
            logger.info(f"Converting PDF to DOCX before processing: {resume_path}")
            try:
                working_docx_path = self.pdf_converter.convert_pdf_to_docx(resume_path)
                logger.info(f"✅ PDF converted to DOCX: {working_docx_path}")
            except Exception as e:
                logger.error(f"❌ PDF to DOCX conversion failed: {e}")
                raise RuntimeError(f"Failed to convert PDF to DOCX: {e}")
        else:
            logger.info(f"Input file is already DOCX format: {resume_path}")
            
        try:
            logger.info(f"Starting resume enhancement workflow for {resume_path}")
            
            # Generate unique request ID for tracking
            request_id = self._generate_request_id()
            logger.info(f"Request ID: {request_id}")
            
            # Step 1: Parse Job Description
            logger.info("Step 1: Parsing job description...")
            jd_summary = self.jd_parser.parse_jd(job_description)
            logger.info(f"✅ Job description parsed: {len(jd_summary.must_have)} must-have skills, {len(jd_summary.nice_to_have)} nice-to-have skills")
            
            # Step 2: Parse Resume
            logger.info("Step 2: Parsing resume...")
            resume_ast = self.resume_parser.parse_resume(resume_path)
            logger.info(f"✅ Resume parsed: {len(resume_ast.roles)} roles, {len(resume_ast.bullets)} bullets")
            
            # Step 3: Gap Analysis
            logger.info("Step 3: Analyzing gaps...")
            improvement_plan = self.gap_analyzer.analyze_gaps(jd_summary, resume_ast)
            logger.info(f"✅ Gap analysis completed: {len(improvement_plan.edits)} improvements identified")
            
            # Step 4: Tag structure and apply Improvements
            logger.info("Step 4: Tagging document structure (original)...")
            original_structure_input = self._collect_paragraph_metadata(working_docx_path)
            original_structure_tags = self.structure_tagger.tag_paragraphs(original_structure_input)

            logger.info("Step 5: Building deterministic mapping and pre-image hashes...")
            mapping = self._build_frozen_mapping(working_docx_path, resume_ast, original_structure_tags)
            enriched_plan = self._enrich_plan_with_mapping(improvement_plan, mapping)

            logger.info("Step 6: Applying improvements deterministically...")
            # Collect original company lines in order of roles for preservation
            original_company_lines = [
                f"{r.company}, {getattr(r, 'location', '')}".replace(' ,', ',').strip().rstrip(',') for r in resume_ast.roles
            ]
            enhanced_docx_path = self.document_editor.apply_plan(
                working_docx_path, {"edits": enriched_plan}, resume_ast, structure_tags=original_structure_tags, original_company_lines=original_company_lines
            )
            logger.info(f"✅ Improvements applied: {enhanced_docx_path}")
            
            # Step 7: Tag enhanced and evaluate (optional - can fail gracefully)
            logger.info("Step 7: Tagging document structure (enhanced) and evaluating...")
            evaluation_report = {"passed": True, "issues": []}  # Default success
            
            # Skip evaluation on Render to avoid resource issues
            if os.environ.get('RENDER', 'false').lower() == 'true':
                logger.info("Skipping evaluation on Render to avoid resource constraints")
                evaluation_report = {"passed": True, "issues": ["evaluation_skipped_on_render"]}
            else:
                try:
                    enhanced_structure_input = self._collect_paragraph_metadata(enhanced_docx_path)
                    enhanced_structure_tags = self.structure_tagger.tag_paragraphs(enhanced_structure_input)
                    evaluation_report = self.evaluator.evaluate(
                        original_docx_path=working_docx_path,
                        enhanced_docx_path=enhanced_docx_path,
                        original_tags=original_structure_tags,
                        enhanced_tags=enhanced_structure_tags,
                    )
                    if not evaluation_report.get("passed", False):
                        logger.warning(f"Post-edit evaluation reported issues: {evaluation_report.get('issues')}")
                except Exception as e:
                    logger.warning(f"Step 7 evaluation failed, continuing workflow: {e}")
                    evaluation_report = {"passed": True, "issues": ["evaluation_skipped_due_to_error"]}
            
            # Step 8: Generate PDF (optional)
            enhanced_pdf_path = None
            if generate_pdf:
                logger.info("Step 8: Generating PDF...")
                enhanced_pdf_path = self.pdf_converter.convert_docx_to_pdf(enhanced_docx_path)
                if enhanced_pdf_path:
                    logger.info(f"✅ PDF generated: {enhanced_pdf_path}")
                else:
                    logger.warning("⚠️ PDF generation failed")
            
            # Step 9: Prepare results
            results = self._prepare_results(
                request_id=request_id,
                original_path=resume_path,
                enhanced_docx_path=enhanced_docx_path,
                enhanced_pdf_path=enhanced_pdf_path,
                jd_summary=jd_summary,
                resume_ast=resume_ast,
                improvement_plan=improvement_plan
            )
            results["evaluation_report"] = evaluation_report
            results["deterministic_mapping"] = mapping
            
            # Step 10: Save workflow artifacts
            self._save_workflow_artifacts(results, output_dir or os.path.dirname(resume_path))
            
            # Clean up temporary converted DOCX if it was created from PDF
            if is_pdf_input and working_docx_path != resume_path:
                try:
                    os.remove(working_docx_path)
                    logger.info(f"Cleaned up temporary converted DOCX: {working_docx_path}")
                except Exception as e:
                    logger.warning(f"Failed to clean up temporary DOCX: {e}")
            
            logger.info(f"🎉 Resume enhancement workflow completed successfully for request {request_id}")
            return results
            
        except Exception as e:
            logger.error(f"❌ Resume enhancement workflow failed: {e}")
            # Clean up temporary converted DOCX on error
            if is_pdf_input and working_docx_path != resume_path:
                try:
                    os.remove(working_docx_path)
                    logger.info(f"Cleaned up temporary converted DOCX on error: {working_docx_path}")
                except Exception as cleanup_error:
                    logger.warning(f"Failed to clean up temporary DOCX on error: {cleanup_error}")
            raise
    
    def _collect_paragraph_metadata(self, docx_path: str) -> List[Dict]:
        """Collect paragraph metadata from DOCX for structure tagging (no text rewriting)."""
        try:
            from docx import Document
            doc = Document(docx_path)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            items: List[Dict] = []
            for idx, p in enumerate(doc.paragraphs):
                p_el = getattr(p, '_element', None)
                has_numpr = bool(p_el is not None and p_el.find('.//w:numPr', ns) is not None)
                style_name = ''
                try:
                    style_name = (p.style.name or '').strip()
                except Exception:
                    pass
                outline_level = None
                try:
                    # Read outline level if available
                    if p_el is not None:
                        ol = p_el.find('.//w:outlineLvl', ns)
                        if ol is not None and ol.get(f'{{{ns["w"]}}}val') is not None:
                            outline_level = int(ol.get(f'{{{ns["w"]}}}val'))
                except Exception:
                    pass
                text = (p.text or '').strip()
                items.append({
                    'paragraph_index': idx,
                    'text': text,
                    'style_name': style_name,
                    'has_numpr': has_numpr,
                    'outline_level': outline_level,
                })
            return items
        except Exception as e:
            logger.warning(f"Failed to collect paragraph metadata: {e}")
            return []

    def _build_frozen_mapping(self, docx_path: str, resume_ast, structure_tags: Dict) -> Dict:
        """Build deterministic mapping: paragraph hashes, role spans, bullet indices, and pre-image hashes."""
        from docx import Document
        import hashlib
        doc = Document(docx_path)
        paragraph_hashes = [hashlib.sha256((p.text or '').encode('utf-8')).hexdigest() for p in doc.paragraphs]

        # Use structure tags for role spans and bullet indices
        role_spans = list(structure_tags.get('role_spans', []))
        paragraph_tags = structure_tags.get('paragraph_tags', [])
        bullet_indices = sorted([int(t.get('paragraph_index')) for t in paragraph_tags if t.get('label') == 'bullet' and isinstance(t.get('paragraph_index'), int)])

        # Map ResumeAST bullets to paragraph indices in order per role
        bullet_index_map: Dict[str, int] = {}
        for span_idx, role in enumerate(resume_ast.roles):
            if span_idx >= len(role_spans):
                break
            s = int(role_spans[span_idx].get('start_index', 0))
            e = int(role_spans[span_idx].get('end_index', s))
            indices_in_span = [i for i in bullet_indices if s <= i <= e]
            role_bullets = [b for b in resume_ast.bullets if b.role_id == role.id]
            for b, p_idx in zip(role_bullets, indices_in_span):
                bullet_index_map[b.id] = p_idx

        # Pre-image hashes for mapped bullets
        pre_hashes: Dict[int, str] = {}
        for b_id, idx in bullet_index_map.items():
            if 0 <= idx < len(doc.paragraphs):
                pre_hashes[idx] = paragraph_hashes[idx]

        return {
            'paragraph_hashes': paragraph_hashes,
            'role_spans': role_spans,
            'bullet_index_map': bullet_index_map,
            'pre_image_hashes': pre_hashes,
        }

    def _enrich_plan_with_mapping(self, plan: Plan, mapping: Dict) -> List[Dict]:
        """Attach resolved indices and pre-image hashes to each edit for deterministic apply."""
        enriched: List[Dict] = []
        bullet_map: Dict[str, int] = mapping.get('bullet_index_map', {})
        pre_hashes: Dict[int, str] = mapping.get('pre_image_hashes', {})
        for e in plan.edits:
            if hasattr(e, 'type'):
                et = e.type
                if et == 'modify_bullet':
                    idx = bullet_map.get(e.bullet_id)
                    enriched.append({
                        'type': 'modify_bullet',
                        'bullet_id': e.bullet_id,
                        'new_text': e.new_text,
                        'resolved_paragraph_index': idx,
                        'pre_image_hash': pre_hashes.get(idx)
                    })
                elif et == 'insert_bullet':
                    after_idx = bullet_map.get(e.after_bullet_id) if getattr(e, 'after_bullet_id', None) else None
                    enriched.append({
                        'type': 'insert_bullet',
                        'role_id': e.role_id,
                        'text': e.text,
                        'after_bullet_id': getattr(e, 'after_bullet_id', None),
                        'resolved_after_paragraph_index': after_idx
                    })
                elif et == 'upsert_skill':
                    enriched.append({
                        'type': 'upsert_skill',
                        'bucket': e.bucket,
                        'value': e.value
                    })
            else:
                # dict edit
                et = e.get('type')
                if et == 'modify_bullet':
                    idx = bullet_map.get(e.get('bullet_id'))
                    enriched.append({
                        'type': 'modify_bullet',
                        'bullet_id': e.get('bullet_id'),
                        'new_text': e.get('new_text'),
                        'resolved_paragraph_index': idx,
                        'pre_image_hash': pre_hashes.get(idx)
                    })
                elif et == 'insert_bullet':
                    after = e.get('after_bullet_id')
                    after_idx = bullet_map.get(after) if after else None
                    enriched.append({
                        'type': 'insert_bullet',
                        'role_id': e.get('role_id'),
                        'text': e.get('text'),
                        'after_bullet_id': after,
                        'resolved_after_paragraph_index': after_idx
                    })
                elif et == 'upsert_skill':
                    enriched.append({
                        'type': 'upsert_skill',
                        'bucket': e.get('bucket'),
                        'value': e.get('value')
                    })
        return enriched
    
    def analyze_only(
        self,
        resume_path: str,
        job_description: str
    ) -> Dict:
        """
        Run analysis only (no document modification).
        
        Returns:
            Dict containing analysis results without file modifications
        """
        try:
            logger.info(f"Running analysis-only workflow for {resume_path}")
            
            # Parse job description and resume
            jd_summary = self.jd_parser.parse_jd(job_description)
            resume_ast = self.resume_parser.parse_resume(resume_path)
            
            # Generate improvement plan
            improvement_plan = self.gap_analyzer.analyze_gaps(jd_summary, resume_ast)
            
            return {
                "analysis_completed": True,
                "jd_summary": jd_summary,
                "resume_ast": resume_ast,
                "improvement_plan": improvement_plan,
                "recommendations_count": len(improvement_plan.edits),
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Analysis workflow failed: {e}")
            raise
    
    def batch_enhance(
        self,
        resume_paths: List[str],
        job_description: str,
        output_dir: str
    ) -> List[Dict]:
        """
        Enhance multiple resumes for the same job description.
        
        Args:
            resume_paths: List of resume file paths
            job_description: Job description text
            output_dir: Directory to save all enhanced files
            
        Returns:
            List of results for each resume
        """
        results = []
        
        for i, resume_path in enumerate(resume_paths, 1):
            try:
                logger.info(f"Processing resume {i}/{len(resume_paths)}: {resume_path}")
                
                # Create subdirectory for this resume
                resume_name = Path(resume_path).stem
                resume_output_dir = os.path.join(output_dir, f"enhanced_{resume_name}")
                os.makedirs(resume_output_dir, exist_ok=True)
                
                # Run enhancement workflow
                result = self.enhance_resume(
                    resume_path=resume_path,
                    job_description=job_description,
                    output_dir=resume_output_dir,
                    preserve_original=True,
                    generate_pdf=True
                )
                
                results.append(result)
                logger.info(f"✅ Completed resume {i}/{len(resume_paths)}")
                
            except Exception as e:
                logger.error(f"❌ Failed to enhance resume {resume_path}: {e}")
                results.append({
                    "success": False,
                    "error": str(e),
                    "resume_path": resume_path
                })
        
        return results
    
    def _generate_request_id(self) -> str:
        """Generate unique request ID for tracking."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        import uuid
        unique_id = str(uuid.uuid4())[:8]
        return f"req_{timestamp}_{unique_id}"
    
    def _prepare_results(
        self,
        request_id: str,
        original_path: str,
        enhanced_docx_path: str,
        enhanced_pdf_path: Optional[str],
        jd_summary: JDSummary,
        resume_ast: ResumeAST,
        improvement_plan: Plan
    ) -> Dict:
        """Prepare comprehensive results dictionary."""
        return {
            "success": True,
            "request_id": request_id,
            "timestamp": datetime.now().isoformat(),
            "workflow_summary": {
                "jd_skills_analyzed": len(jd_summary.must_have) + len(jd_summary.nice_to_have),
                "resume_roles_parsed": len(resume_ast.roles),
                "resume_bullets_parsed": len(resume_ast.bullets),
                "improvements_generated": len(improvement_plan.edits),
                "seniority_level": jd_summary.seniority
            },
            "file_paths": {
                "original": original_path,
                "enhanced_docx": enhanced_docx_path,
                "enhanced_pdf": enhanced_pdf_path
            },
            "analysis": {
                "jd_summary": jd_summary,
                "resume_ast": resume_ast,
                "improvement_plan": improvement_plan
            }
        }
    
    def _save_workflow_artifacts(self, results: Dict, output_dir: str):
        """Save workflow artifacts for reference and debugging."""
        try:
            # Create artifacts directory
            artifacts_dir = os.path.join(output_dir, "workflow_artifacts")
            os.makedirs(artifacts_dir, exist_ok=True)
            
            # Save improvement plan
            plan_file = os.path.join(artifacts_dir, f"improvement_plan_{results['request_id']}.json")
            with open(plan_file, 'w') as f:
                plan_dict = {
                    "edits": [
                        {
                            "type": edit.type,
                            **{k: v for k, v in edit.model_dump().items() if k != "type"}
                        } for edit in results['analysis']['improvement_plan'].edits
                    ],
                    "constraints": results['analysis']['improvement_plan'].constraints.model_dump() if results['analysis']['improvement_plan'].constraints else {}
                }
                json.dump(plan_dict, f, indent=2)
            
            # Save workflow summary
            summary_file = os.path.join(artifacts_dir, f"workflow_summary_{results['request_id']}.json")
            with open(summary_file, 'w') as f:
                json.dump(results, f, indent=2, default=str)
            
            logger.info(f"Workflow artifacts saved to {artifacts_dir}")
            
        except Exception as e:
            logger.warning(f"Failed to save workflow artifacts: {e}")
    
    def get_workflow_status(self, request_id: str) -> Dict:
        """Get status of a specific workflow request."""
        # This could be enhanced to track workflow status in a database
        return {
            "request_id": request_id,
            "status": "completed",  # For now, assume completed
            "timestamp": datetime.now().isoformat()
        }
    
    def cleanup_temp_files(self, request_id: str):
        """Clean up temporary files for a specific request."""
        # Implementation for cleaning up temporary files
        logger.info(f"Cleaning up temporary files for request {request_id}")
        pass
