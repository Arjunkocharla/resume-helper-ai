from docx import Document
from docx.shared import Inches
from .models import EnhancementSuggestion
import logging
from typing import Dict, List

class DocumentEnhancer:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def enhance_document(self, doc: Document, suggestions: Dict[str, List[EnhancementSuggestion]]) -> Document:
        """Apply enhancements to document"""
        section_map = self._create_section_map(doc)
        
        for section_name, section_suggestions in suggestions.items():
            if section_name in section_map:
                self._apply_section_suggestions(
                    doc,
                    section_map[section_name],
                    section_suggestions
                )
        
        return doc
    
    def _create_section_map(self, doc: Document) -> dict:
        """Create mapping of sections to their locations"""
        section_map = {}
        current_section = None
        
        common_section_keywords = {
            'experience': ['experience', 'work ex', 'employment', 'work history'],
            'education': ['education', 'academic', 'qualifications'],
            'skills': ['skills', 'technical skills', 'expertise'],
            'projects': ['projects', 'personal projects'],
            'profile': ['profile', 'summary', 'objective']
        }
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text.isupper():  # Potential section header
                # Store exact section name as it appears
                current_section = text
                section_map[current_section] = {
                    'start': i,
                    'end': None,
                    'type': self._get_section_type(text, common_section_keywords)
                }
            elif current_section:
                section_map[current_section]['end'] = i
                
        return section_map

    def _get_section_type(self, text: str, keywords: dict) -> str:
        """Helper to identify section type while preserving original name"""
        text_lower = text.lower()
        for section_type, variations in keywords.items():
            if any(kw in text_lower for kw in variations):
                return section_type
        return 'other'

    def _apply_section_suggestions(self, doc: Document, section_loc: Dict, suggestions: List[EnhancementSuggestion]):
        """Apply suggestions to a specific section"""
        start_idx = section_loc['start']
        end_idx = section_loc['end']

        current_role = None
        role_start_idx = None

        # Process each paragraph in the section
        for i in range(start_idx + 1, end_idx if end_idx else len(doc.paragraphs)):
            para_text = doc.paragraphs[i].text.strip()
            
            # Check if this is a role header (not a bullet point)
            if para_text and not para_text.startswith('•'):
                if current_role:
                    # Add suggestions for previous role
                    self._add_role_suggestions(doc, role_start_idx, i, suggestions, current_role)
                current_role = para_text
                role_start_idx = i
        
        # Add suggestions for the last role
        if current_role:
            self._add_role_suggestions(doc, role_start_idx, end_idx, suggestions, current_role)
        # Add suggestions without roles (like Skills section)
        elif suggestions:
            self._add_general_suggestions(doc, start_idx, end_idx, suggestions)

    def _add_role_suggestions(self, doc: Document, start_idx: int, end_idx: int, 
                            suggestions: List[EnhancementSuggestion], role: str):
        """Add suggestions under a specific role"""
        role_suggestions = [s for s in suggestions if s.role == role]
        for suggestion in role_suggestions:
            new_para = doc.add_paragraph(suggestion.suggestion)
            new_para.paragraph_format.left_indent = Inches(0.5)
            doc.paragraphs[end_idx-1]._p.addnext(new_para._p)

    def _add_general_suggestions(self, doc: Document, start_idx: int, end_idx: int, 
                               suggestions: List[EnhancementSuggestion]):
        """Add suggestions to sections without roles (like Skills)"""
        for suggestion in suggestions:
            new_para = doc.add_paragraph(suggestion.suggestion)
            new_para.paragraph_format.left_indent = Inches(0.5)
            if end_idx:
                doc.paragraphs[end_idx-1]._p.addnext(new_para._p)
            else:
                doc.paragraphs[start_idx]._p.addnext(new_para._p)
