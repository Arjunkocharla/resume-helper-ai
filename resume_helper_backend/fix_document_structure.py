#!/usr/bin/env python3
"""Quick fix for document structure issues - preserving company lines and removing duplicates."""

import os
from docx import Document

def fix_document_structure(docx_path: str) -> str:
    """Fix common document structure issues after editing."""
    doc = Document(docx_path)
    
    # Track changes made
    changes_made = []
    
    # Fix 1: Ensure TCS company line is present after "Software Development Engineer"
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "Software Development Engineer" in text and i < len(doc.paragraphs) - 1:
            next_text = doc.paragraphs[i + 1].text.strip()
            # If next line is a bullet or doesn't contain company info, insert TCS line
            if (next_text.startswith("•") or 
                ("Tata Consultancy Services" not in next_text and "Hyderabad" not in next_text)):
                
                # Insert TCS company line
                new_para = doc.add_paragraph()
                new_para.text = "Tata Consultancy Services, Hyderabad, India"
                
                # Move to correct position (after the role header)
                doc._body._body.insert(i + 1, new_para._element)
                changes_made.append(f"Inserted TCS company line at position {i + 1}")
                break
    
    # Fix 2: Remove duplicate bullets (keep the one that's properly positioned)
    seen_bullets = set()
    to_remove = []
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("•"):
            # Normalize bullet text for comparison
            normalized = text.lower().replace("•", "").strip()
            if normalized in seen_bullets:
                # This is a duplicate - mark for removal
                to_remove.append(i)
                changes_made.append(f"Marked duplicate bullet for removal: {text[:50]}...")
            else:
                seen_bullets.add(normalized)
    
    # Remove duplicates (from bottom up to maintain indices)
    for i in reversed(to_remove):
        elem = doc.paragraphs[i]._element
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
    
    # Fix 3: Remove orphaned bullets that appear before proper section headers
    orphaned_to_remove = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("•"):
            # Check if this bullet is orphaned (no role context within 5 lines before)
            has_role_context = False
            for j in range(max(0, i-5), i):
                prev_text = doc.paragraphs[j].text.strip()
                if any(keyword in prev_text for keyword in [
                    "Engineer", "Developer", "Intern", "Tutor", "Manager", 
                    "CentrAlert", "CAMP Systems", "University", "Tata Consultancy"
                ]):
                    has_role_context = True
                    break
            
            if not has_role_context:
                orphaned_to_remove.append(i)
                changes_made.append(f"Marked orphaned bullet for removal: {text[:50]}...")
    
    # Remove orphaned bullets
    for i in reversed(orphaned_to_remove):
        elem = doc.paragraphs[i]._element
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
    
    # Save fixed document
    fixed_path = docx_path.replace(".docx", "_fixed.docx")
    doc.save(fixed_path)
    
    print(f"🔧 Document structure fixes applied:")
    for change in changes_made:
        print(f"  - {change}")
    
    return fixed_path

if __name__ == "__main__":
    # Fix the latest enhanced document
    latest_doc = "nagarjuna_kocharla_resume_main_enhanced_21.docx"
    if os.path.exists(latest_doc):
        fixed_doc = fix_document_structure(latest_doc)
        print(f"✅ Fixed document saved as: {fixed_doc}")
    else:
        print(f"❌ Document not found: {latest_doc}")
