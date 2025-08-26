#!/usr/bin/env python3
"""Debug script to examine document structure and understand why roles aren't detected."""

import os
import sys
from docx import Document

def debug_document_structure():
    """Debug the document structure to understand role detection issues."""
    
    resume_path = "nagarjuna_kocharla_resume_main.docx"
    if not os.path.exists(resume_path):
        print(f"❌ Resume file not found: {resume_path}")
        return
    
    doc = Document(resume_path)
    
    print("🔍 DEBUGGING DOCUMENT STRUCTURE")
    print("=" * 60)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print()
    
    # Examine each paragraph
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        print(f"Paragraph {i:2d}: {text}")
        
        # Check if this looks like a role
        if ' at ' in text:
            print(f"        ⭐ CONTAINS 'at' - potential role!")
            
            # Check for company indicators
            company_indicators = [
                'inc', 'corp', 'llc', 'ltd', 'company', 'tech', 'systems', 
                'consultancy', 'international', 'solutions', 'group', 'partners',
                'associates', 'enterprises', 'ventures', 'holdings', 'industries'
            ]
            
            found_indicators = []
            for keyword in company_indicators:
                if keyword in text.lower():
                    found_indicators.append(keyword)
            
            if found_indicators:
                print(f"        ✅ Company indicators found: {found_indicators}")
            else:
                print(f"        ❌ No company indicators found")
            
            # Check for date patterns
            import re
            date_patterns = [
                r'\d{4}\s*-\s*\d{4}',  # 2019-2021
                r'\d{2}/\d{2}\s*-\s*Present',  # 06/23 - Present
                r'\d{4}\s*-\s*Present',  # 2019 - Present
                r'\d{2}/\d{2}\s*-\s*\d{2}/\d{2}',  # 06/19 - 07/21
            ]
            
            found_dates = []
            for pattern in date_patterns:
                if re.search(pattern, text):
                    found_dates.append(re.search(pattern, text).group())
            
            if found_dates:
                print(f"        ✅ Date patterns found: {found_dates}")
            else:
                print(f"        ❌ No date patterns found")
        
        # Check for job title patterns
        job_titles = [
            'engineer', 'developer', 'architect', 'manager', 'lead', 'consultant',
            'analyst', 'specialist', 'coordinator', 'director', 'officer'
        ]
        
        found_titles = []
        for title in job_titles:
            if title.lower() in text.lower():
                found_titles.append(title)
        
        if found_titles and ' at ' in text:
            print(f"        🎯 Job titles found: {found_titles}")
        
        print()

if __name__ == "__main__":
    debug_document_structure()
