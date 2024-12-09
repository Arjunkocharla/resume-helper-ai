import os
from flask import Flask, request, jsonify, send_file
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document
from flask_cors import CORS 
import json
import anthropic
import re
from b2sdk.v2 import InMemoryAccountInfo, B2Api
import io
import tempfile
import time
from dotenv import load_dotenv
import boto3
from botocore.client import Config
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PyPDF2 import PdfReader, PdfWriter
import openai  # Import OpenAI for GPT
import logging
from flask import Flask, request, jsonify, send_file
import os
import re
import json
from werkzeug.utils import secure_filename
from docx import Document
import tempfile
import anthropic  # Assuming you're using the Anthropic client library
import PyPDF2  # For PDF text extraction
from typing import Dict, Optional, Union
from docx import Document
import json
import os
from werkzeug.utils import secure_filename
import logging
from anthropic import Anthropic
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import difflib
from copy import deepcopy
from docx.shared import Pt, Inches
from docx import Document
import logging
import os
import time
from werkzeug.utils import secure_filename
from flask import jsonify, send_file
from typing import Dict, Optional, Tuple
import re

# Load environment variables
load_dotenv()
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
app = Flask(__name__)
CORS(app,origins='*')
#client = anthropic.Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY'))

# Replace the current initialization
# client = anthropic.Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY'))

# With this:
client = anthropic.Anthropic(
    api_key=os.getenv('ANTHROPIC_API_KEY'),
    base_url="https://api.anthropic.com",
)

gpt_client = openai.OpenAI(
    api_key = os.getenv("GPT_API_KEY"),
)


# Configure upload folder and allowed extensions
UPLOAD_FOLDER = os.path.abspath(os.path.join(os.path.dirname(__file__), 'uploads'))
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Backblaze B2 configuration
B2_KEY_ID = os.getenv('B2_KEY_ID')
B2_APPLICATION_KEY = os.getenv('B2_APPLICATION_KEY')
B2_BUCKET_NAME = os.getenv('B2_BUCKET_NAME')
B2_ENDPOINT = os.getenv('B2_ENDPOINT')

# Initialize B2 API
info = InMemoryAccountInfo()
b2_api = B2Api(info)
b2_api.authorize_account("production", B2_KEY_ID, B2_APPLICATION_KEY)
bucket = b2_api.get_bucket_by_name(B2_BUCKET_NAME)

# S3 client for generating pre-signed URLs
s3_client = boto3.client(
    's3',
    endpoint_url=B2_ENDPOINT,
    aws_access_key_id=B2_KEY_ID,
    aws_secret_access_key=B2_APPLICATION_KEY,
    config=Config(signature_version='s3v4')
)

# Initialize OpenAI API
openai.api_key = os.getenv('GPT_API_KEY')

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(file_path):
    if file_path.endswith('.pdf'):
        return extract_pdf_text(file_path)
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    else:
        raise ValueError('Unsupported file format')

def analyze_keywords_with_claude(resume_text, job_category):
    prompt = f"""As an expert HR specialist with deep knowledge of the {job_category} industry, analyze this resume for a {job_category} position.

First, identify the main sections in the resume. Then, for each suggested improvement, specify which of these EXACT sections it should be placed under.

Resume text for section analysis:
{resume_text}

Now, identify the top 8-9 most relevant keywords or phrases for this specific {job_category} role, considering current industry trends and job market demands.
For each keyword:
   a. Provide a brief explanation of its importance in the {job_category} field (1-2 sentences).
   b. Generate 1-2 impactful, ready-to-use bullet points that the candidate could directly add to their resume. These points should:
      - Integrate the identified keyword naturally
      - Be tailored specifically to the {job_category} position
      - Highlight quantifiable achievements where possible
      - Use strong action verbs
      - Demonstrate the candidate's potential impact and value
   c. Specify which EXACT section from the resume the bullet points should be added to.

Provide the output in this JSON format:
{{
  "sections": [
    "EDUCATION",
    "SKILLS & TOOLS",
    "COMPUTER SCIENCE PROJECTS",
    "WORK EXPERIENCE",
    // ... exact sections as they appear in the resume
  ],
  "keywords": [
    {{
      "keyword": "string",
      "importance": "string",
      "bullet_points": [
        {{
          "point": "string",
          "section": "EXACT_SECTION_NAME"  // Must match one of the sections listed above
        }}
      ]
    }}
  ]
}}"""

    response = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=2000,
        temperature=0.2,
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    # Extract the JSON part from the response
    response_content = response.content[0].text
    json_start = response_content.find('{')
    json_end = response_content.rfind('}') + 1
    json_str = response_content[json_start:json_end]

    return json_str

@app.route('/analyze_resume_structure', methods=['POST'])
def analyze_resume_structure():
    if 'resume' not in request.files:
        return jsonify({'error': 'Resume file is required.'}), 400

    file = request.files['resume']
    retry = request.form.get('retry', 'false').lower() == 'true'

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file. Only PDF and DOCX are allowed.'}), 400

    # Ensure the upload folder exists
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    resume_text = extract_text(file_path)

    retry_prefix = """IMPORTANT: Previous response had JSON parsing errors. Please provide ONLY the JSON response in the exact format specified below. No additional text, explanations or casual conversation - just the parseable JSON response.

""" if retry else ""

    prompt = f"""{retry_prefix}As a senior resume analyst and ATS expert with extensive experience in multiple industries, conduct a comprehensive analysis of the following resume. Focus on ATS compatibility, keyword optimization, and industry-specific best practices.
    Respond ONLY with valid JSON in the exact format shown below and is parseable with json loads, with no additional text or explanations.
    1. "ATS_Compatibility_Score": An integer from 1-100 assessing how well the resume would perform in ATS scans. Consider factors such as formatting, use of standard section headings, and keyword relevance.

    2. "Keywords_Analysis": {{
        "Present_Keywords": An array of all relevant keywords found in the resume, sorted by importance.
        "Missing_Keywords": An array of important industry-standard keywords not present in the resume, based on the inferred job category.
        "Keyword_Density": A string representing the overall keyword density (e.g., "15.3%").
        "Keyword_Distribution": An object showing the distribution of keywords across different resume sections.
    }}

    3. "ATS_Friendly_Structure": {{
        "Format_Score": An integer from 1-10 assessing how ATS-friendly the resume format is.
        "Section_Order": An array of strings representing the current order of resume sections.
        "Recommended_Section_Order": An array of strings suggesting the ideal ATS-friendly section order.
        "Formatting_Issues": An array of specific formatting issues that could hinder ATS parsing.
    }}

    4. "Content_Analysis": {{
        "Total_Word_Count": An integer.
        "Bullet_Points_Count": An integer.
        "Action_Verbs_Count": An integer.
        "Quantifiable_Achievements_Count": An integer.
        "Average_Bullet_Length": A float representing the average number of words per bullet point.
        "Skills_Section_Analysis": An object analyzing the effectiveness of the skills section.
    }}

    5. "ATS_Optimization_Tips": An array of strings with specific, actionable recommendations to improve ATS compatibility and keyword usage. Prioritize the top 5 most impactful changes.

    6. "Strengths": An array of strings highlighting the resume's strong points, focusing on elements that would appeal to both ATS systems and human recruiters.

    7. "Industry_Specific_Suggestions": {{
        "Inferred_Industry": A string indicating the primary industry or job category this resume targets.
        "Industry_Keywords": An array of 20-30 highly relevant industry-specific keywords to consider incorporating, ranked by importance.
        "Industry_Specific_Tips": An array of strings with industry-specific advice for ATS optimization and standing out in the inferred field.
        "Emerging_Trends": An array of 3-5 emerging trends or skills in the inferred industry that could enhance the resume.
    }}

    8. "Overall_Assessment": A string (3-4 sentences) evaluating the resume's effectiveness for ATS and human review, highlighting major strengths and areas for improvement. Provide a clear, actionable next step for the candidate.

    9. "Industry_Standards": {{
        "Word_Count_Range": A string representing the typical word count range for resumes in this industry (e.g., "400-600").
        "Keyword_Density_Range": A string representing the ideal keyword density range (e.g., "2%-3.5%").
        "Bullet_Points_Range": A string representing the ideal number of bullet points (e.g., "15-25").
        "Sections_Importance": An object with section names as keys and importance scores (1-10) as values, tailored to the inferred industry.
        "Key_Action_Verbs": An array of 10-15 powerful action verbs commonly used in top resumes for this industry.
        "Ideal_Quantifiable_Achievements": An integer representing the ideal number of quantifiable achievements.
        "File_Format_Preference": A string indicating the preferred file format(s) for ATS (e.g., "PDF, DOCX").
        "Optimal_Length_Pages": A string representing the ideal resume length in pages (e.g., "1-2").
    }}

    10. "Tailored_Improvement_Plan": An array of 3-5 specific, prioritized steps the candidate should take to improve their resume's effectiveness, based on the analysis.

    Resume Text:
    {resume_text}

    Ensure all output is ATS-friendly, industry-specific, and focused on maximizing the resume's performance in both automated screening systems and human review. The JSON should be valid and parseable without errors. Do not include any text outside of the JSON structure.
    """

    try:
        response = client.messages.create(
            model="claude-3-haiku-20240307",
            max_tokens=3000,
            temperature=0.2,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        # Extract the content from the response
        response_content = response.content[0].text

        # Log the response content for debugging
        print("Response Content:", response_content)

        # Find the JSON part of the response
        json_start = response_content.find('{')
        json_end = response_content.rfind('}') + 1
        
        if json_start == -1 or json_end == 0:
            # If no JSON found, return a fallback response
            return jsonify({
                'ATS_Compatibility_Score': 50,
                'Keywords_Analysis': {
                    'Present_Keywords': [],
                    'Missing_Keywords': [],
                    'Keyword_Density': '0%',
                    'Keyword_Distribution': {}
                },
                'ATS_Friendly_Structure': {
                    'Format_Score': 5,
                    'Section_Order': [],
                    'Recommended_Section_Order': ['Contact Information', 'Summary', 'Experience', 'Education', 'Skills'],
                    'Formatting_Issues': ['Unable to analyze resume format']
                },
                'Content_Analysis': {
                    'Total_Word_Count': len(resume_text.split()),
                    'Bullet_Points_Count': 0,
                    'Action_Verbs_Count': 0,
                    'Quantifiable_Achievements_Count': 0,
                    'Average_Bullet_Length': 0.0,
                    'Skills_Section_Analysis': {}
                },
                'ATS_Optimization_Tips': ['Unable to generate specific recommendations'],
                'Strengths': ['Unable to identify specific strengths'],
                'Industry_Specific_Suggestions': {
                    'Inferred_Industry': 'Unknown',
                    'Industry_Keywords': [],
                    'Industry_Specific_Tips': [],
                    'Emerging_Trends': []
                },
                'Overall_Assessment': 'Unable to perform detailed analysis. Please ensure the resume is in a readable format.',
                'Industry_Standards': {
                    'Word_Count_Range': '400-600',
                    'Keyword_Density_Range': '2%-3.5%',
                    'Bullet_Points_Range': '15-25',
                    'Sections_Importance': {},
                    'Key_Action_Verbs': [],
                    'Ideal_Quantifiable_Achievements': 5,
                    'File_Format_Preference': 'PDF, DOCX',
                    'Optimal_Length_Pages': '1-2'
                },
                'Tailored_Improvement_Plan': ['Review resume format', 'Ensure content is readable', 'Try submitting again']
            }), 200

        # Extract and try to parse the JSON part
        json_str = response_content[json_start:json_end]
        try:
            analysis = json.loads(json_str)
            return jsonify(analysis), 200
        except json.JSONDecodeError:
            # If JSON is invalid, try to clean it up
            cleaned_json = json_str.replace('\n', ' ').replace('\r', '')
            try:
                analysis = json.loads(cleaned_json)
                return jsonify(analysis), 200
            except json.JSONDecodeError:
                # If still invalid, return the fallback response
                return jsonify({
                    'error': 'Failed to parse analysis results',
                    'message': 'The resume analysis service is temporarily unavailable. Please try again.'
                }), 500

    except Exception as e:
        return jsonify({
            'error': str(e),
            'message': 'An unexpected error occurred during analysis.'
        }), 500

@app.route('/analyze_keywords', methods=['POST'])
def analyze_keywords():
    if 'resume' not in request.files or 'job_category' not in request.form:
        return jsonify({'error': 'Resume file and job category are required.'}), 400

    file = request.files['resume']
    job_category = request.form['job_category']

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file. Only PDF and DOCX are allowed.'}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    resume_text = extract_text(file_path)

    try:
        analysis_json = analyze_keywords_with_claude(resume_text, job_category)
        analysis = json.loads(analysis_json)  # Parse the JSON string here
        return jsonify({
            'job_category': job_category,
            'analysis': analysis
        }), 200

    except json.JSONDecodeError as e:
        return jsonify({'error': f'Failed to parse JSON: {str(e)}'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/analyze_resume_length', methods=['POST'])
def analyze_resume_length():
    if 'resume' not in request.files or 'job_category' not in request.form:
        return jsonify({'error': 'Resume file and job category are required.'}), 400

    file = request.files['resume']
    job_category = request.form['job_category']

    if file.filename == '':
        return jsonify({'error': 'No selected file.'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Extract text from resume
        resume_text = extract_text(file_path)
        word_count = len(resume_text.split())

        # Crafting the prompt
        prompt = f"""As an expert career advisor specializing in the {job_category} industry, analyze the following resume text and determine if the length is optimal for making a strong impression in this specific field. Consider current industry standards and expectations for the {job_category} sector.

        1. Provide a detailed analysis of the resume length, including:
           - The current word count
           - The ideal word count range for {job_category} resumes
           - Whether the resume is too short, too long, or just right
           - The impact of the current length on ATS scans and human readability

        2. If adjustments are needed, suggest specific ways to improve the resume length, such as:
           - Sections to expand or condense
           - Types of information to add or remove
           - Strategies for concise yet impactful writing in the {job_category} field

        3. Explain how the optimal length contributes to the overall effectiveness of the resume for {job_category} positions.

        4. Provide 2-3 industry-specific tips for balancing comprehensiveness with conciseness in {job_category} resumes.

        Resume Text:
        {resume_text}

        Format your response as a JSON object with the following structure:
        {{
          "current_word_count": integer,
          "ideal_word_count_range": "string",
          "length_assessment": "string",
          "ats_impact": "string",
          "human_readability_impact": "string",
          "improvement_suggestions": [
            "string",
            ...
          ],
          "length_importance_explanation": "string",
          "industry_specific_tips": [
            "string",
            ...
          ]
        }}
        """

        try:
            response = client.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=1000,
                temperature=0.3,
                messages=[{"role": "user", "content": prompt}
                ]
            )

            # Extract the JSON part from the response
            response_content = response.content[0].text
            json_start = response_content.find('{')
            json_end = response_content.rfind('}') + 1
            analysis = json.loads(response_content[json_start:json_end])

            return jsonify(analysis), 200

        except json.JSONDecodeError as e:
            return jsonify({'error': f'Failed to parse JSON: {str(e)}'}), 500
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    else:
        return jsonify({'error': 'Invalid file type. Only PDF and DOCX are allowed.'}), 400

@app.route('/suggest_keywords', methods=['POST'])
def suggest_keywords():
    if 'resume' not in request.files or 'job_description' not in request.form:
        return jsonify({'error': 'Resume file and job description are required.'}), 400

    file = request.files['resume']
    job_description = request.form['job_description']
    retry = request.form.get('retry', 'false').lower() == 'true'

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file. Only PDF and DOCX are allowed.'}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    resume_text = extract_text(file_path)

    retry_prefix = """IMPORTANT: Previous response had JSON parsing errors. Please provide ONLY the JSON response in the exact format specified below. No additional text, explanations or casual conversation - just the parseable JSON response.

""" if retry else ""

    prompt = f"""{retry_prefix}As an expert ATS optimization specialist and industry recruiter, perform a deep analysis of this resume and job description to provide highly specific, tailored 6-7 keyword suggestions. Focus on actionable, industry-specific improvements that will maximize ATS scoring.
    Respond ONLY with valid JSON in the exact format shown below and is parseable with json loads, with no additional text or explanations.

1. First, analyze the resume to understand:
   - The candidate's current experience level and role
   - Technical skills, tools, and domain expertise
   - Industry-specific achievements and certifications
   - Current role and career trajectory
   - Existing keyword density and placement

2. Then, analyze the job description to identify:
   - Must-have technical skills, tools, and technologies
   - Required certifications and qualifications
   - Desired experience levels and competencies
   - Industry-specific terminology and frameworks
   - Key responsibilities and deliverables
   - Recurring keywords and their variations
   - Technology stack requirements
   - Domain-specific methodologies

3. Based on this analysis, provide:
   - High-impact keywords missing from the resume but crucial for ATS scoring
   - Existing keywords that need stronger emphasis or modern context
   - Industry-specific technical terms that would strengthen the application
   - Role-specific tools and technologies mentioned in the job description
   - Methodologies and frameworks valued in the industry
   - Measurable metrics and achievements using these keywords

For each keyword suggestion:
- Explain why it's specifically important for this role and industry
- Detail how it impacts ATS scoring and ranking
- Provide 2-3 ready-to-use bullet points that:
  * Incorporate the keyword naturally and with proper context
  * Include specific metrics and quantifiable achievements
  * Use strong action verbs aligned with seniority level
  * Are tailored to the candidate's experience level
  * Follow ATS-friendly formatting and keyword placement
  * Are industry-specific and technically accurate
  * Include relevant tools and methodologies
  * Demonstrate impact and results

Resume Text:
{resume_text}

Job Description:
{job_description}

Provide your response in the following JSON format:
{{
  "experience_gap_analysis": "A detailed analysis of the gap between current resume and job requirements",
  "keywords": [
    {{
      "keyword": "Specific technical or professional term",
      "importance": "Detailed explanation of why this keyword is crucial for this specific role and its impact on ATS scoring",
      "bullet_points": [
        {{
          "point": "Complete, ready-to-use bullet point with metrics and context",
          "explanation": "Why this bullet point is effective and how it strengthens the resume for ATS optimization",
          "under_experience": "Specific company/role under which this bullet point should be added (e.g., 'Software Engineer at Google', 'Full Stack Developer at Microsoft')"
        }}
      ],
      "placement": "Specific section where this keyword/bullet point should be added for maximum ATS impact"
    }}
  ],
  "overall_strategy": "Comprehensive strategy for implementing these changes effectively and optimizing ATS scoring"
}}
"""

    try:
        response = client.messages.create(
            model="claude-3-haiku-20240307",
            max_tokens=2000,
            temperature=0.3,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )

        response_content = response.content[0].text
        json_start = response_content.find('{')
        json_end = response_content.rfind('}') + 1
        json_str = response_content[json_start:json_end]

        # Sanitize the JSON string
        json_str = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', json_str)

        # Add error handling for JSON parsing
        try:
            suggestions = json.loads(json_str)
        except json.JSONDecodeError as json_error:
            # If JSON parsing fails, return the error details
            return jsonify({
                'error': f'Failed to parse JSON: {str(json_error)}',
                'json_str': json_str,
                'error_position': json_error.pos,
                'error_lineno': json_error.lineno,
                'error_colno': json_error.colno
            }), 500

        return jsonify({
            'job_description': job_description,
            'keyword_suggestions': suggestions
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/upload_resume', methods=['POST'])
def upload_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'Resume file is required.'}), 400

    file = request.files['resume']

    if file.filename == '':
        return jsonify({'error': 'No selected file.'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        
        try:
            # Initialize B2 API
            info = InMemoryAccountInfo()
            b2_api = B2Api(info)
            b2_api.authorize_account("production", B2_KEY_ID, B2_APPLICATION_KEY)

            # Get the bucket
            bucket = b2_api.get_bucket_by_name(B2_BUCKET_NAME)

            # Upload the file
            file_data = file.read()
            file_info = bucket.upload_bytes(
                data_bytes=file_data,
                file_name=filename,
                content_type=file.content_type
            )

            # Verify the upload
            if not file_info or not file_info.id_:
                raise Exception("File upload failed: No file info returned")

            # Get the file URL
            download_url = b2_api.get_download_url_for_fileid(file_info.id_)

            # Verify the download URL
            if not download_url:
                raise Exception("Failed to generate download URL")

            # Use b2_api to get file info instead of bucket
            file_metadata = b2_api.get_file_info(file_info.id_)

            return jsonify({
                'message': 'File uploaded successfully',
                'filename': filename,
                'file_id': file_info.id_,
                'download_url': download_url,
                'file_size': file_metadata.size if file_metadata else None
            }), 200

        except Exception as e:
            app.logger.error(f"Upload error: {str(e)}")
            return jsonify({'error': f'Upload failed: {str(e)}'}), 500

    else:
        return jsonify({'error': 'Invalid file type. Only PDF, DOC, and DOCX are allowed.'}), 400

import os
from flask import send_file, request, jsonify
from b2sdk.v2 import InMemoryAccountInfo, B2Api
from b2sdk.v2.exception import B2Error
import io

@app.route('/download_resume', methods=['GET'])
def download_resume():
    file_id = request.args.get('file_id')
    
    if not file_id:
        return jsonify({'error': 'file_id is required'}), 400

    try:
        # Initialize B2 API
        info = InMemoryAccountInfo()
        b2_api = B2Api(info)
        b2_api.authorize_account("production", B2_KEY_ID, B2_APPLICATION_KEY)

        # Get file info
        file_info = b2_api.get_file_info(file_id)

        if not file_info:
            return jsonify({'error': 'File not found'}), 404

        # Get the bucket
        bucket = b2_api.get_bucket_by_name(B2_BUCKET_NAME)

        # Download the file to a BytesIO object
        download_buffer = io.BytesIO()
        bucket.download_file_by_id(file_id, download_buffer)
        download_buffer.seek(0)

        # Send the file
        return send_file(
            download_buffer,
            as_attachment=True,
            download_name=file_info.file_name,
            mimetype=file_info.content_type
        )

    except B2Error as e:
        app.logger.error(f"B2 API Error: {str(e)}")
        return jsonify({
            'error': f'B2 API Error: {str(e)}',
            'file_id': file_id
        }), 500
    except Exception as e:
        app.logger.error(f"Download error: {str(e)}")
        return jsonify({
            'error': f'Download failed: {str(e)}',
            'file_id': file_id
        }), 500

@app.route('/get_resume_download_url', methods=['GET'])
def get_resume_download_url():
    file_id = request.args.get('file_id')
    
    if not file_id:
        return jsonify({'error': 'file_id is required'}), 400

    try:
        # Initialize B2 API to get file info
        info = InMemoryAccountInfo()
        b2_api = B2Api(info)
        b2_api.authorize_account("production", B2_KEY_ID, B2_APPLICATION_KEY)

        # Get file info
        file_info = b2_api.get_file_info(file_id)

        if not file_info:
            return jsonify({'error': 'File not found'}), 404

        # Use boto3 to generate pre-signed URL
        s3_client = boto3.client(
            's3',
            endpoint_url=B2_ENDPOINT,  # Use the endpoint from environment variables
            aws_access_key_id=B2_KEY_ID,
            aws_secret_access_key=B2_APPLICATION_KEY,
            config=Config(signature_version='s3v4')
        )

        duration_in_seconds = 3600  # 1 hour

        # Generate the pre-signed URL
        presigned_url = s3_client.generate_presigned_url(
            'get_object',
            Params={
                'Bucket': B2_BUCKET_NAME,
                'Key': file_info.file_name
            },
            ExpiresIn=duration_in_seconds
        )

        return jsonify({
            'download_url': presigned_url,
            'filename': file_info.file_name,
            'content_type': file_info.content_type,
            'size': file_info.size,
            'expires_in': duration_in_seconds
        }), 200

    except Exception as e:
        app.logger.error(f"Error generating pre-signed URL: {str(e)}")
        return jsonify({
            'error': f'Failed to generate pre-signed URL: {str(e)}',
            'file_id': file_id
        }), 500

def update_resume_via_gpt(resume_text, suggestions):
    messages = [
        {"role": "system", "content": "You are a resume editor."},
        {"role": "user", "content": f"Here is a resume and some suggestions for improvement. Please apply the suggestions to the resume.\n\nResume Text:\n{resume_text}\n\nSuggestions:\n{suggestions}\n\nReturn the updated resume text."}
    ]
    response = gpt_client.chat.completions.create(
        model="gpt-4",  # Use the appropriate model
        messages=messages,
        max_tokens=2000,
        temperature=0.3
    )
    return response.choices[0].message.content.strip()

def create_pdf_with_text(text, output_path):
    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    can.drawString(10, 750, text)  # Adjust the position as needed
    can.save()

    packet.seek(0)
    with open(output_path, 'wb') as f:
        f.write(packet.getbuffer())

@app.route('/apply_suggestions', methods=['POST'])
def apply_suggestions():
    if 'resume' not in request.files or 'suggestions' not in request.form:
        return jsonify({'error': 'Resume file and suggestions are required.'}), 400

    file = request.files['resume']
    suggestions = json.loads(request.form['suggestions'])

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file. Only PDF and DOCX are allowed.'}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    resume_text = extract_text(file_path)

    # Generate suggestions using Claude
    prompt = f"""As an expert resume editor, analyze the following resume and suggestions. 
    Provide a list of specific changes to be made to the resume, including the exact text to be added, 
    removed, or modified, and the location of each change. Do not rewrite the entire resume.

    Resume Text:
    {resume_text}

    Suggestions:
    {json.dumps(suggestions, indent=2)}

    Please return your response in the following JSON format, that is parsable and json loads can handle:
    {{
        "changes": [
            {{
                "type": "add" | "remove" | "modify",
                "location": "section name or line number",
                "original_text": "text to be changed (for modify and remove)",
                "new_text": "text to be added or modified"
            }}
        ]
    }}
    """
    try:
        response = client.messages.create(
            model="claude-3-haiku-20240307",
            max_tokens=2000,
            temperature=0.3,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        # Extract the suggestions from the Claude model response
        suggestions_text = response.content[0].text

        # Pass the resume text and suggestions to GPT for updates
        updated_resume_text = update_resume_via_gpt(resume_text, suggestions_text)

        # Create a new PDF with the updated resume text
        updated_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"updated_{filename}")
        create_pdf_with_text(updated_resume_text, updated_file_path)

        # Upload the updated file to Backblaze B2
        download_url = upload_to_backblaze(updated_file_path)

        return jsonify({'download_url': download_url}), 200

    except Exception as e:
        logging.error(f"Error in apply_suggestions: {e}")
        return jsonify({'error': str(e)}), 500

def upload_to_backblaze(file_path):
    # Initialize B2 API
    info = InMemoryAccountInfo()
    b2_api = B2Api(info)
    b2_api.authorize_account("production", B2_KEY_ID, B2_APPLICATION_KEY)

    # Get the bucket
    bucket = b2_api.get_bucket_by_name(B2_BUCKET_NAME)

    # Upload the file
    with open(file_path, 'rb') as file_data:
        file_info = bucket.upload_bytes(
            data_bytes=file_data.read(),
            file_name=os.path.basename(file_path),
            content_type='application/pdf'  # or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' for DOCX
        )

    # Get the download URL
    download_url = b2_api.get_download_url_for_fileid(file_info.id_)
    return download_url

def apply_changes_to_file(file_path, changes):
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return apply_changes_to_pdf(file_path, changes)
    elif file_extension == '.docx':
        return apply_changes_to_docx(file_path, changes)
    else:
        raise ValueError("Unsupported file type")

def apply_changes_to_pdf(file_path, changes):
    # Read the existing PDF
    reader = PdfReader(file_path)
    output = PdfWriter()

    # Iterate over each page
    for page_number, page in enumerate(reader.pages):
        # Extract text from the page
        text = page.extract_text()

        # Apply changes
        for change in changes:
            if change['type'] == 'modify':
                text = text.replace(change['original_text'], change['new_text'])
            elif change['type'] == 'add':
                text += f"\n{change['new_text']}"
            elif change['type'] == 'remove':
                text = text.replace(change['original_text'], '')

        # Create a new PDF page with the modified text
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        can.drawString(10, 750, text)  # Adjust the position as needed
        can.save()

        # Move to the beginning of the StringIO buffer
        packet.seek(0)
        new_pdf = PdfReader(packet)
        output.add_page(new_pdf.pages[0])

    # Save the updated PDF
    updated_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"updated_{os.path.basename(file_path)}")
    with open(updated_file_path, 'wb') as f:
        output.write(f)

    return updated_file_path

def apply_changes_to_docx(file_path, changes):
    doc = Document(file_path)
    
    for change in changes:
        if change['type'] == 'modify':
            for paragraph in doc.paragraphs:
                if change['original_text'] in paragraph.text:
                    paragraph.text = paragraph.text.replace(change['original_text'], change['new_text'])
        elif change['type'] == 'add':
            doc.add_paragraph(change['new_text'])
        elif change['type'] == 'remove':
            for paragraph in doc.paragraphs:
                if change['original_text'] in paragraph.text:
                    paragraph.text = paragraph.text.replace(change['original_text'], '')
    
    updated_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"updated_{os.path.basename(file_path)}")
    doc.save(updated_file_path)
    
    return updated_file_path

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(file_path):
    # Implement text extraction for both DOCX and PDF
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif file_path.endswith('.pdf'):
        # Implement PDF text extraction
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ''
            for page in reader.pages:
                text += page.extract_text() + '\n'
        return text
    else:
        return ''

def identify_sections_with_llm(resume_text):
    prompt = f"""As an AI language model, please analyze the following resume text and extract only the section headings exactly as they appear in the document. Return only the exact section names without any additional text or explanations.

Resume text:
{resume_text}

Return only a JSON object in this exact format:
{{
    "sections": [
        "Section Name 1",
        "Section Name 2",
        "Section Name 3"
    ]
}}

Important Instructions:
- Extract ONLY the section headings from the resume.
- Do NOT include any additional text, explanations, or descriptions.
- Do NOT modify the original section names.
- Do NOT add words like 'section', 'under the', etc.
- Do NOT include any text other than the JSON object in the specified format.
"""

    try:
        response = client.messages.create(
            model="claude-3-haiku-20240307",
            max_tokens=1000,
            temperature=0.2,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        response_content = response.content[0].text
        json_start = response_content.find('{')
        json_end = response_content.rfind('}') + 1
        sections_data = json.loads(response_content[json_start:json_end])
        return sections_data['sections']
    except Exception as e:
        print(f"Error identifying sections: {e}")
        return None
    
def clean_section_name(section_name: str) -> str:
    """Remove extra words and normalize section names"""
    # Remove common phrases
    phrases_to_remove = [
        'under the',
        'in the',
        'section',
        'sections',
        'under',
        'in',
        'for',
        'the'
    ]
    
    cleaned = section_name.lower().strip()
    for phrase in phrases_to_remove:
        cleaned = cleaned.replace(phrase, '').strip()
    
    return cleaned

def update_resume_docx(doc_path: str, suggestions: dict) -> str:
    try:
        doc = Document(doc_path)
        resume_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Keep existing section mappings and print statements
        section_mappings = {
            'work experience': ['work experience', 'experience', 'employment', 'work history'],
            'projects': ['projects', 'project experience', 'technical projects'],
            'internships': ['internships', 'internship experience'],
            'education': ['education', 'academic background'],
            'skills': ['skills', 'technical skills', 'skills & tools'],
            'profile': ['profile', 'summary', 'professional summary']
        }
        
        # Keep existing sections mapping code and print statements
        sections = {}
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.lower().strip()
            for section_key, variations in section_mappings.items():
                if any(var in para_text for var in variations):
                    sections[section_key] = {
                        'index': i,
                        'text': paragraph.text,
                        'roles': {}  # Add roles dictionary to store role information
                    }
                    print(f"Found section: {section_key} at index {i}")

        print("\n=== IDENTIFIED SECTIONS ===")
        print(json.dumps(sections, indent=2))

        # Add role identification within work experience section
        current_section = None
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Identify current section
            if text and text.isupper():
                current_section = text
                print(f"Current section: {current_section}")
            
            # Within work experience, identify roles by looking for lines with dates
            if current_section == "WORK EXPERIENCE" and ":" in text and not text.startswith('•'):
                role_name = text.split(":")[0].strip()
                print(f"Found role: {role_name} at index {i}")
                if 'work experience' in sections:
                    sections['work experience']['roles'][role_name] = {
                        'index': i,
                        'text': text
                    }

        # Modify bullet point processing to use under_experience
        for keyword in suggestions.get('keywords', []):
            for bullet in keyword.get('bullet_points', []):
                if isinstance(bullet, dict) and 'point' in bullet and 'under_experience' in bullet:
                    target_role = bullet['under_experience']
                    print(f"\nProcessing bullet point for role: {target_role}")
                    print(f"Bullet point: {bullet['point']}")
                    
                    # Try to find matching role
                    role_found = False
                    if 'work experience' in sections:
                        for role_name, role_info in sections['work experience']['roles'].items():
                            if any(part.lower() in role_name.lower() for part in target_role.lower().split()):
                                print(f"Found matching role: {role_name}")
                                insert_index = role_info['index'] + 1
                                
                                # Move past existing bullets
                                while (insert_index < len(doc.paragraphs) and 
                                       doc.paragraphs[insert_index].text.strip().startswith('•')):
                                    insert_index += 1
                                    print(f"Skipping existing bullet at index {insert_index}")
                                
                                print(f"Adding bullet point at index {insert_index}")
                                p = doc.paragraphs[insert_index-1].insert_paragraph_before("• " + bullet['point'])
                                if insert_index < len(doc.paragraphs):
                                    p.style = doc.paragraphs[insert_index].style
                                role_found = True
                                break
                    
                    if not role_found:
                        print(f"Warning: Could not find matching role for '{target_role}'")

        # Keep existing save and return code
        output_dir = os.path.dirname(doc_path)
        updated_path = os.path.join(output_dir, f"updated_{os.path.basename(doc_path)}")
        doc.save(updated_path)
        return updated_path

    except Exception as e:
        print(f"Error in update_resume_docx: {e}")
        raise

def update_resume_pdf(pdf_path, suggestions):
    # For PDFs, consider converting to DOCX, updating, then exporting back to PDF
    # This is a complex process and may require third-party tools or services
    # For simplicity, we'll return an error for now
    return None

@app.route('/update_resumes', methods=['POST'])
def update_resume():
    if 'resume' not in request.files or 'job_description' not in request.form:
        return jsonify({'error': 'Resume file and job description are required.'}), 400

    file = request.files['resume']
    job_description = request.form['job_description']
    retry = request.form.get('retry', 'false').lower() == 'true'

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file. Only PDF and DOCX are allowed.'}), 400

    filename = secure_filename(file.filename)
    file_ext = filename.rsplit('.', 1)[1].lower()

    # Create a directory for processed files
    processed_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'processed')
    os.makedirs(processed_dir, exist_ok=True)

    try:
        # Save the file in the processed directory
        file_path = os.path.join(processed_dir, filename)
        file.save(file_path)

        resume_text = extract_text(file_path)

        retry_prefix = """IMPORTANT: Previous response had JSON parsing errors. Please provide ONLY the JSON response in the exact format specified below. No additional text, explanations or casual conversation - just the parseable JSON response.

""" if retry else ""

        prompt = f"""{retry_prefix}As an expert ATS optimization specialist and industry recruiter, perform a deep analysis of this resume and job description to provide highly specific, tailored 6-7 keyword suggestions. Focus on actionable, industry-specific improvements that will maximize ATS scoring.
Respond ONLY with valid JSON in the exact format shown below and is parseable with json loads, with no additional text or explanations.

1. First, analyze the resume to understand:
   - The candidate's current experience level and role
   - Technical skills, tools, and domain expertise
   - Industry-specific achievements and certifications
   - Current role and career trajectory
   - Existing keyword density and placement

2. Then, analyze the job description to identify:
   - Must-have technical skills, tools, and technologies
   - Required certifications and qualifications
   - Desired experience levels and competencies
   - Industry-specific terminology and frameworks
   - Key responsibilities and deliverables
   - Recurring keywords and their variations
   - Technology stack requirements
   - Domain-specific methodologies

3. Based on this analysis, provide:
   - High-impact keywords missing from the resume but crucial for ATS scoring
   - Existing keywords that need stronger emphasis or modern context
   - Industry-specific technical terms that would strengthen the application
   - Role-specific tools and technologies mentioned in the job description
   - Methodologies and frameworks valued in the industry
   - Measurable metrics and achievements using these keywords

For each keyword suggestion:
- Explain why it's specifically important for this role and industry
- Detail how it impacts ATS scoring and ranking
- Provide 2-3 ready-to-use bullet points that:
  * Incorporate the keyword naturally and with proper context
  * Include specific metrics and quantifiable achievements
  * Use strong action verbs aligned with seniority level
  * Are tailored to the candidate's experience level
  * Follow ATS-friendly formatting and keyword placement
  * Are industry-specific and technically accurate
  * Include relevant tools and methodologies
  * Demonstrate impact and results

Resume Text:
{resume_text}

Job Description:
{job_description}

Provide your response in the following JSON format:
{{
  "experience_gap_analysis": "A detailed analysis of the gap between current resume and job requirements",
  "keywords": [
    {{
      "keyword": "Specific technical or professional term",
      "importance": "Detailed explanation of why this keyword is crucial for this specific role and its impact on ATS scoring",
      "bullet_points": [
        {{
          "point": "Complete, ready-to-use bullet point with metrics and context",
          "explanation": "Why this bullet point is effective and how it strengthens the resume for ATS optimization",
          "under_experience": "Specific company/role under which this bullet point should be added (e.g., 'Software Engineer at Google', 'Full Stack Developer at Microsoft')"
        }}
      ],
      "placement": "Specific section where this keyword/bullet point should be added for maximum ATS impact"
    }}
  ],
  "overall_strategy": "Comprehensive strategy for implementing these changes effectively and optimizing ATS scoring"
}}
"""

        try:
            response = client.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=2000,
                temperature=0.3,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )

            response_content = response.content[0].text
            json_start = response_content.find('{')
            json_end = response_content.rfind('}') + 1
            json_str = response_content[json_start:json_end]
            json_str = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', json_str)

            try:
                suggestions = json.loads(json_str)
            except json.JSONDecodeError as json_error:
                return jsonify({
                    'error': f'Failed to parse JSON: {str(json_error)}',
                    'json_str': json_str,
                    'error_position': json_error.pos,
                    'error_lineno': json_error.lineno,
                    'error_colno': json_error.colno
                }), 500

            # Update the resume
            if file_ext == 'docx':
                updated_resume_path = update_resume_docx(file_path, suggestions)
            elif file_ext == 'pdf':
                updated_resume_path = update_resume_pdf(file_path, suggestions)
                if not updated_resume_path:
                    return jsonify({'error': 'PDF resume updating is not supported at this time.'}), 500
            else:
                return jsonify({'error': 'Unsupported file format.'}), 400

            if not os.path.exists(updated_resume_path):
                return jsonify({'error': 'Updated resume file not found.'}), 500

            try:
                # Create a downloads directory in the current working directory
                downloads_dir = os.path.join(os.getcwd(), 'downloads')
                os.makedirs(downloads_dir, exist_ok=True)

                # Save the file in the downloads directory with a timestamp
                timestamp = time.strftime("%Y%m%d-%H%M%S")
                download_filename = f"updated_{timestamp}_{filename}"
                download_path = os.path.join(downloads_dir, download_filename)

                try:
                    # Instead of sending the file directly, copy it to downloads directory
                    import shutil
                    shutil.copy2(updated_resume_path, download_path)

                    # Send both the file and the saved location
                    response = send_file(
                        updated_resume_path,
                        as_attachment=True,
                        download_name=download_filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_ext == 'docx' else 'application/pdf'
                    )

                    # Add the saved path to response headers
                    response.headers['X-Saved-Location'] = download_path

                    @response.call_on_close
                    def cleanup():
                        try:
                            # Only remove the temporary files, keep the downloaded copy
                            if os.path.exists(file_path):
                                os.remove(file_path)
                            if os.path.exists(updated_resume_path):
                                os.remove(updated_resume_path)
                        except Exception as e:
                            print(f"Cleanup error: {e}")

                    return response

                except Exception as e:
                    return jsonify({
                        'error': f'Error sending file: {str(e)}',
                        'saved_location': download_path if os.path.exists(download_path) else None
                    }), 500

            except Exception as e:
                return jsonify({'error': str(e)}), 500

        except Exception as e:
            return jsonify({'error': str(e)}), 500

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # Modified cleanup to preserve the downloaded copy
        try:
            if 'file_path' in locals() and os.path.exists(file_path):
                os.remove(file_path)
            # Don't remove updated_resume_path as it's been copied to downloads
        except Exception as e:
            print(f"Final cleanup error: {e}")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResumeUpdateResult:
    def __init__(self, 
                 success: bool, 
                 file_path: Optional[str] = None, 
                 error: Optional[str] = None,
                 changes_made: Optional[Dict] = None):
        self.success = success
        self.file_path = file_path
        self.error = error
        self.changes_made = changes_made or {}

class ResumeUpdater:
    def __init__(self, client: Anthropic):
        self.client = client
        
    def _extract_text(self, file_path: str) -> str:
        """Extract text from DOCX file while preserving formatting"""
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return '\n'.join(text)

    def _create_update_prompt(self, resume_text: str, suggestions: Dict) -> str:
        """Create a detailed prompt for the LLM"""
        return f"""As an expert resume editor, you will update this resume by applying the suggested changes.

IMPORTANT INSTRUCTIONS:
1. Return ONLY the complete updated resume text.
2. Preserve the exact formatting, including:
   - Section headers in UPPERCASE
   - Bullet points using '•'
   - Original spacing and line breaks
3. Add new bullet points under the EXACT role/company specified
4. Maintain consistent formatting with existing bullet points
5. Do not modify or remove existing content unless explicitly instructed
6. Ensure all new bullet points start with '•'

Original Resume:
{resume_text}

Suggested Updates:
{json.dumps(suggestions, indent=2)}

Return only the complete updated resume text, preserving all formatting and structure."""

    def _validate_updates(self, original_text: str, updated_text: str) -> Dict:
        """Validate that updates were applied correctly"""
        changes = {
            'added_lines': [],
            'removed_lines': [],
            'modified_lines': []
        }
        
        # Use difflib to find differences
        diff = difflib.unified_diff(
            original_text.splitlines(),
            updated_text.splitlines(),
            lineterm=''
        )
        
        for line in diff:
            if line.startswith('+•'):
                changes['added_lines'].append(line[1:])
            elif line.startswith('-•'):
                changes['removed_lines'].append(line[1:])
            elif line.startswith('+') or line.startswith('-'):
                changes['modified_lines'].append(line[1:])
                
        return changes

    def _format_document(self, doc: Document) -> None:
        """Apply consistent formatting to the document"""
        for para in doc.paragraphs:
            if para.text.strip().isupper() and len(para.text.split()) <= 4:
                # Section headers
                run = para.runs[0] if para.runs else para.add_run()
                run.font.bold = True
                run.font.size = Pt(12)
            elif para.text.strip().startswith('•'):
                # Bullet points
                run = para.runs[0] if para.runs else para.add_run()
                run.font.size = Pt(11)

    def update_resume(self, file_path: str, suggestions: Dict) -> ResumeUpdateResult:
        """Update resume using LLM and return the result"""
        try:
            # Extract original text
            original_text = self._extract_text(file_path)
            
            # Create prompt and get LLM response
            prompt = self._create_update_prompt(original_text, suggestions)
            
            response = self.client.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=2000,
                temperature=0.2,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            updated_text = response.content[0].text.strip()
            
            # Validate updates
            changes = self._validate_updates(original_text, updated_text)
            
            if not changes['added_lines'] and not changes['modified_lines']:
                return ResumeUpdateResult(
                    success=False,
                    error="No changes were applied to the resume"
                )
            
            # Create new document with updates
            output_path = os.path.join(
                os.path.dirname(file_path),
                f"updated_{os.path.basename(file_path)}"
            )
            
            new_doc = Document()
            for line in updated_text.split('\n'):
                new_doc.add_paragraph(line)
            
            # Apply consistent formatting
            self._format_document(new_doc)
            
            # Save updated document
            new_doc.save(output_path)
            
            return ResumeUpdateResult(
                success=True,
                file_path=output_path,
                changes_made=changes
            )
            
        except Exception as e:
            logger.error(f"Error updating resume: {str(e)}")
            return ResumeUpdateResult(
                success=False,
                error=str(e)
            )

@app.route('/update_resume_enhanced', methods=['POST'])
def update_resume_enhanced():
    if 'resume' not in request.files or 'job_description' not in request.form:
        return jsonify({'error': 'Resume file and job description are required.'}), 400

    file = request.files['resume']
    job_description = request.form['job_description']

    try:
        suggestions = json.loads(request.form['suggestions'])
    except json.JSONDecodeError:
        return jsonify({'error': 'Invalid suggestions format'}), 400

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file. Only DOCX files are supported.'}), 400

    try:
        # Save the uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Initialize resume updater with Anthropic client
        updater = ResumeUpdater(client)
        
        # Update the resume
        result = updater.update_resume(file_path, suggestions)
        
        if not result.success:
            return jsonify({
                'error': result.error or 'Failed to update resume'
            }), 500

        # Create a downloads directory in the current working directory
        downloads_dir = os.path.join(os.getcwd(), 'downloads')
        os.makedirs(downloads_dir, exist_ok=True)

        # Save the file in the downloads directory with a timestamp
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        download_filename = f"updated_{timestamp}_{filename}"
        download_path = os.path.join(downloads_dir, download_filename)

        # Copy the updated file to the downloads directory
        import shutil
        shutil.copy2(result.file_path, download_path)

        # Send the updated file
        response = send_file(
            result.file_path,
            as_attachment=True,
            download_name=f"updated_{filename}",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Add changes to response headers
        response.headers['X-Changes-Made'] = json.dumps(result.changes_made)
        response.headers['X-Saved-Location'] = download_path

        return response

    except Exception as e:
        logger.error(f"Error in update_resume_enhanced: {e}")
        return jsonify({'error': str(e)}), 500

    finally:
        # Cleanup temporary files
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
            if 'result' in locals() and result.file_path and os.path.exists(result.file_path):
                os.remove(result.file_path)
        except Exception as e:
            logger.error(f"Cleanup error: {e}")

@app.route('/update_resume_with_suggestions', methods=['POST'])
def update_resume_with_suggestions():
    # Validate input files
    if 'resume' not in request.files or 'description' not in request.form:
        return jsonify({'error': 'Resume file and job description are required.'}), 400

    file = request.files['resume']
    job_description = request.form['description']

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file. Only PDF and DOCX are allowed.'}), 400

    try:
        # Save the uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Extract text from resume
        resume_text = extract_text(file_path)

        # Single prompt to analyze and update resume
        prompt = f"""As an expert ATS optimization specialist and resume editor, analyze this job description and update the resume to improve ATS scoring while maintaining exact formatting.

Job Description:
{job_description}

Resume:
{resume_text}

Instructions:
1. Analyze the job description for key skills, qualifications, and terminology
2. Update the resume by adding relevant bullet points and skills
3. Preserve ALL existing content - do not remove anything
4. Maintain exact formatting including:
   - All bullet points using '•' symbol
   - All section headers in UPPERCASE
   - All spacing and line breaks
   - Contact information format
   - Date formats
   - Company/role formats

Return only the complete updated resume text with all formatting preserved."""

        # Get updated resume from Claude
        response = client.messages.create(
            model="claude-3-haiku-20240307",
            max_tokens=2000,
            temperature=0.2,
            messages=[{"role": "user", "content": prompt}]
        )

        updated_resume_text = response.content[0].text.strip()

        # Create downloads directory
        downloads_dir = os.path.join(os.getcwd(), 'downloads')
        os.makedirs(downloads_dir, exist_ok=True)

        # Save with timestamp
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        download_filename = f"updated_{timestamp}_{filename}"
        download_path = os.path.join(downloads_dir, download_filename)

        # Create new document and save updated text
        new_doc = Document()
        for line in updated_resume_text.split('\n'):
            new_doc.add_paragraph(line)
        new_doc.save(download_path)

        # Send response
        return send_file(
            download_path,
            as_attachment=True,
            download_name=download_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"Error in update_resume_with_suggestions: {str(e)}")
        return jsonify({'error': str(e)}), 500

    finally:
        # Cleanup temporary files
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            logger.error(f"Cleanup error: {e}")

class InPlaceResumeUpdater:
    def __init__(self, original_doc_path: str):
        self.doc = Document(original_doc_path)
        self.sections = self._map_sections()
        
    def _map_sections(self) -> Dict[str, Tuple[int, int]]:
        """Maps section names to their start and end indices in the document"""
        sections = {}
        current_section = None
        section_start = 0
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if text.isupper() and len(text.split()) <= 4:
                # End previous section
                if current_section:
                    sections[current_section] = (section_start, i - 1)
                # Start new section
                current_section = text
                section_start = i
                
        # Add final section
        if current_section:
            sections[current_section] = (section_start, len(self.doc.paragraphs) - 1)
            
        return sections
    
    def update_resume(self, updated_content: str, output_path: str) -> None:
        """Updates resume content while preserving formatting"""
        # Parse updated content into sections
        updated_sections = {}
        current_section = None
        current_content = []
        
        for line in updated_content.split('\n'):
            line = line.strip()
            if line.isupper() and len(line.split()) <= 4:
                # Save previous section
                if current_section:
                    updated_sections[current_section] = current_content
                # Start new section
                current_section = line
                current_content = []
            elif line:
                current_content.append(line)
                
        # Save final section
        if current_section:
            updated_sections[current_section] = current_content
            
        # Create new document based on original
        new_doc = Document()
        for style in self.doc.styles:
            if style.name not in new_doc.styles:
                new_doc.styles.add_style(style.name, style.type, True)
        
        # Update each section
        for section_name, (start_idx, end_idx) in self.sections.items():
            if section_name in updated_sections:
                # Get formatting templates from existing paragraphs
                templates = {
                    'bullet': None,
                    'normal': None
                }
                
                # Remove existing content after section header
                for i in range(start_idx + 1, end_idx + 1):
                    para = self.doc.paragraphs[i]
                    if para.text.strip().startswith('•') and not templates['bullet']:
                        templates['bullet'] = para
                    elif not para.text.strip().startswith('•') and not templates['normal']:
                        templates['normal'] = para
                    if templates['bullet'] and templates['normal']:
                        break
                
                # Add updated content with preserved formatting
                for line in updated_sections[section_name]:
                    template = templates['bullet'] if line.startswith('•') else templates['normal']
                    if template:
                        new_para = new_doc.add_paragraph()
                        new_para._p.append(deepcopy(template._p))
                        new_para.text = line
                    else:
                        # Fallback if no template found
                        new_para = new_doc.add_paragraph(line)
                        
                    # Ensure bullet point formatting
                    if line.startswith('•'):
                        new_para.paragraph_format.left_indent = Inches(0.25)
                        new_para.paragraph_format.first_line_indent = Inches(-0.25)
            else:
                # If section not in updated content, copy original content
                for i in range(start_idx + 1, end_idx + 1):
                    para = self.doc.paragraphs[i]
                    new_para = new_doc.add_paragraph()
                    new_para._p.append(deepcopy(para._p))
        
        # Save the updated document
        new_doc.save(output_path)

@app.route('/inplace_resume_update', methods=['POST'])
def inplace_resume_update():
    """Endpoint for updating resume while preserving formatting"""
    if 'resume' not in request.files or 'description' not in request.form:
        return jsonify({'error': 'Resume file and job description are required.'}), 400

    file = request.files['resume']
    job_description = request.form['description']

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file. Only DOCX files are allowed.'}), 400

    try:
        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Extract text while preserving structure
        doc = Document(file_path)
        resume_text = '\n'.join([para.text for para in doc.paragraphs])

        # Get updated content from Claude
        prompt = f"""As an expert ATS optimization specialist and resume editor, analyze this job description and update the resume to improve ATS scoring while maintaining the resume's current writing style and tone.

Job Description:
{job_description}

Resume:
{resume_text}

Instructions:
1. Analyze the job description for key skills, qualifications, and terminology
2. Update the resume by adding relevant bullet points and skills/tools that:
   - Match the resume's existing writing style and tone for skills/tools sections (e.g. if skills are listed as single words, use single words)
   - Demonstrate concrete impact and value for the target role
   - Use action verbs and quantifiable metrics where appropriate
   - Do not duplicate any existing skills/tools
3. Preserve ALL existing content - do not remove anything
4. Maintain exact formatting including:
   - All bullet points using '•' symbol 
   - All section headers in UPPERCASE (e.g. SKILLS, TOOLS, TECHNICAL SKILLS etc.)
   - All spacing and line breaks
   - Contact information format
   - Date formats
   - Company/role formats
   - Skills/tools formatting and categorization
5. Ensure each added bullet point:
   - Follows the resume's current writing style
   - Clearly demonstrates relevant skills/experience for the role
   - Sounds natural and flows with surrounding content
   - Shows specific impact and value delivered
6. For skills/tools sections:
   - Match the exact format used (e.g. comma-separated, bullet points, categories)
   - Add only relevant missing skills/tools from job description
   - Do not duplicate any existing skills/tools
   - Maintain any existing categorization (e.g. Languages, Frameworks, Tools)

Return only the complete updated resume text with all formatting preserved."""

        response = client.messages.create(
            model="claude-3-haiku-20240307",
            max_tokens=2000,
            temperature=0.2,
            messages=[{"role": "user", "content": prompt}]
        )

        updated_content = response.content[0].text.strip()

        # Create downloads directory
        downloads_dir = os.path.join(os.getcwd(), 'downloads')
        os.makedirs(downloads_dir, exist_ok=True)

        # Save with timestamp
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        download_filename = f"updated_{timestamp}_{filename}"
        download_path = os.path.join(downloads_dir, download_filename)

        # Update resume in-place
        updater = InPlaceResumeUpdater(file_path)
        updater.update_resume(updated_content, download_path)

        return send_file(
            download_path,
            as_attachment=True,
            download_name=download_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"Error in inplace_resume_update: {str(e)}")
        return jsonify({'error': str(e)}), 500

    finally:
        # Cleanup temporary files
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            logger.error(f"Cleanup error: {e}")

import logging
import os
import time
import json
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, Inches
from werkzeug.utils import secure_filename
from flask import jsonify, send_file
from copy import deepcopy

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

class ResumeSection:
    def __init__(self, name: str, start_idx: int):
        self.name = name
        self.start_idx = start_idx
        self.end_idx: Optional[int] = None
        self.roles: Dict[str, Tuple[int, int]] = {}

class StructuredResumeUpdater:
    def __init__(self, file_path: str):
        try:
            self.doc = Document(file_path)
            self.sections = self._map_document_structure()
            logger.info(f"Loaded document with sections: {[s.name for s in self.sections.values()]}")
        except Exception as e:
            logger.error(f"Failed to initialize document: {str(e)}")
            raise

    def _map_document_structure(self) -> Dict[str, ResumeSection]:
        """Maps the document structure including sections and roles"""
        sections = {}
        current_section = None
        current_role = None
        role_start_idx = None

        try:
            for i, para in enumerate(self.doc.paragraphs):
                text = para.text.strip()
                
                # Skip empty paragraphs
                if not text:
                    continue

                # Identify section headers
                if text.isupper() and len(text.split()) <= 4:
                    logger.debug(f"Found section header: {text} at index {i}")
                    
                    # Close previous section
                    if current_section:
                        current_section.end_idx = i - 1
                    
                    # Start new section
                    current_section = ResumeSection(text, i)
                    sections[text] = current_section
                    current_role = None
                
                # Identify roles in work experience
                elif current_section and current_section.name == "WORK EXPERIENCE":
                    # Look for role patterns (e.g., "Company Name: Role" or date patterns)
                    if ":" in text and not text.startswith('•'):
                        if current_role:
                            # Close previous role
                            current_section.roles[current_role] = (role_start_idx, i - 1)
                        
                        current_role = text.split(":")[0].strip()
                        role_start_idx = i
                        logger.debug(f"Found role: {current_role} at index {i}")

            # Close final section and role
            if current_section:
                current_section.end_idx = len(self.doc.paragraphs) - 1
                if current_role:
                    current_section.roles[current_role] = (role_start_idx, len(self.doc.paragraphs) - 1)

            return sections

        except Exception as e:
            logger.error(f"Error mapping document structure: {str(e)}")
            raise

    def _copy_paragraph_formatting(self, source_para, new_para):
        """Copies formatting from source paragraph to new paragraph"""
        try:
            # Copy paragraph format
            if source_para._element.pPr is not None:
                new_para._element.get_or_add_pPr().append(
                    deepcopy(source_para._element.pPr)
                )
            
            # Copy run format
            if source_para.runs:
                new_run = new_para.runs[0] if new_para.runs else new_para.add_run()
                source_run = source_para.runs[0]
                new_run.font.name = source_run.font.name
                new_run.font.size = source_run.font.size
                new_run.font.bold = source_run.font.bold
                new_run.font.italic = source_run.font.italic
        
        except Exception as e:
            logger.warning(f"Error copying paragraph formatting: {str(e)}")

    def update_resume(self, suggestions: Dict[str, List[Dict]], output_path: str) -> None:
        """Updates the resume with the provided suggestions"""
        try:
            new_doc = Document()
            
            # Copy all styles from original document
            for style in self.doc.styles:
                if style.name not in new_doc.styles:
                    new_doc.styles.add_style(style.name, style.type, True)

            current_para_idx = 0
            
            # Process each paragraph
            while current_para_idx < len(self.doc.paragraphs):
                para = self.doc.paragraphs[current_para_idx]
                text = para.text.strip()
                
                # Copy current paragraph
                new_para = new_doc.add_paragraph()
                new_para._p.append(deepcopy(para._p))
                
                # Check if this is a section header
                if text.isupper() and len(text.split()) <= 4:
                    section = self.sections.get(text)
                    if section and text in suggestions:
                        # Process suggestions for this section
                        for suggestion in suggestions[text]:
                            role = suggestion.get('role')
                            content = suggestion.get('suggestion')
                            
                            if role:
                                # Find the role and add suggestion after it
                                for role_name, (start, end) in section.roles.items():
                                    if role.lower() in role_name.lower():
                                        # Skip to role position
                                        while current_para_idx < start:
                                            current_para_idx += 1
                                            new_para = new_doc.add_paragraph()
                                            new_para._p.append(
                                                deepcopy(self.doc.paragraphs[current_para_idx]._p)
                                            )
                                        
                                        # Add suggestion
                                        suggestion_para = new_doc.add_paragraph(content)
                                        self._copy_paragraph_formatting(
                                            self.doc.paragraphs[start + 1],
                                            suggestion_para
                                        )
                            else:
                                # Add suggestion at the end of the section
                                suggestion_para = new_doc.add_paragraph(content)
                                self._copy_paragraph_formatting(para, suggestion_para)
                
                current_para_idx += 1
            
            # Save the updated document
            new_doc.save(output_path)
            logger.info(f"Successfully saved updated document to: {output_path}")
            
        except Exception as e:
            logger.error(f"Error updating resume: {str(e)}")
            raise

@app.route('/structured_resume_update', methods=['POST'])
def structured_resume_update():
    if 'resume' not in request.files or 'description' not in request.form:
        return jsonify({'error': 'Resume file and job description are required.'}), 400

    file = request.files['resume']
    job_description = request.form['description']

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file. Only DOCX files are allowed.'}), 400

    try:
        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Extract text while preserving structure
        doc = Document(file_path)
        resume_text = '\n'.join([para.text for para in doc.paragraphs])

        prompt = f"""As an expert ATS optimization specialist and resume editor, analyze this job description and update the resume to improve ATS scoring while maintaining exact formatting.

Job Description:
{job_description}

Resume:
{resume_text}

Instructions:
1. Analyze the job description for key skills, qualifications, and terminology.
2. Provide suggestions for each relevant section of the resume.
3. For each suggestion, specify:
   - The section it belongs to (EXACTLY as it appears in the resume, in UPPERCASE)
   - The specific role it pertains to (if applicable)
   - The suggested bullet point or content
   - Make sure for each role or project has atleast one suggestion
   - Each point should be unique and not a duplicate of any other point
4. If user mentioned skill in single word, use single word in suggestion as well.
5. Your goal to to improve the resume for ATS scoring and return it in the same format as the resume, for example if user has skills divided into categories like languages, frameworks, tools, etc.
6. Always remember, it defeats the purpose if you rewrite the same point or skill etc in the resume, in your previous response, i saw you did that, please dont do that.
7. Always remember, you cant take a point from one role or project and use it in another role or project, in your previous response, i saw you did that, please dont do that.

Return ONLY a JSON object in this exact format:
{{
    "WORK EXPERIENCE": [
        {{
            "role": "Software Engineer at Google",
            "suggestion": "• Implemented distributed caching system improving response times by 40%"
        }}
    ],
    "SKILLS": [
        {{
            "suggestion": "•Python, Django, FastAPI"
        }}
    ]
    this is just an example, dont copy the content from it
}}"""
        # response = client.messages.create(
        #     model="claude-3-haiku-20240307",
        #     max_tokens=2000,
        #     temperature=0.2,
        #     messages=[{"role": "user", "content": prompt}]
        # )

        # # Extract and parse JSON from response
        # response_content = response.content[0].text

        response = gpt_client.chat.completions.create(
            model="gpt-3.5-turbo",
            temperature=0.2,
            messages=[{"role": "user", "content": prompt}]
        )

        # Extract and parse JSON from response
        response_content = response.choices[0].message.content
        json_start = response_content.find('{')
        json_end = response_content.rfind('}') + 1
        suggestions = json.loads(response_content[json_start:json_end])
        print(suggestions)

        # Apply suggestions in place
        for section, updates in suggestions.items():
            logger.info(f"Processing section: {section}")
            section_found = False
            
            # Find the section and its index
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip().upper() == section:
                    section_found = True
                    logger.info(f"Found section '{section}' at index {i}")
                    
                    # Find section end
                    section_end = len(doc.paragraphs)
                    for j in range(i + 1, len(doc.paragraphs)):
                        if doc.paragraphs[j].text.strip().isupper() and len(doc.paragraphs[j].text.strip().split()) <= 4:
                            section_end = j
                            break
                    
                    # Handle section based on its type
                    section_type = get_section_type(section)
                    
                    if section_type == 'skills':
                        # Look for any category-like lines (ending with colon)
                        for update in updates:
                            suggestion_text = update.get('suggestion', '')
                            if suggestion_text:
                                # Try to find a relevant category, or add to end of skills section
                                category_found = False
                                for k in range(i + 1, section_end):
                                    if doc.paragraphs[k].text.strip().endswith(':'):
                                        doc.paragraphs[k]._p.addnext(doc.add_paragraph(suggestion_text)._p)
                                        category_found = True
                                        logger.info(f"Added skill under category: {doc.paragraphs[k].text.strip()}")
                                        break
                                
                                if not category_found:
                                    # Add to end of skills section if no categories found
                                    doc.paragraphs[section_end-1]._p.addnext(doc.add_paragraph(suggestion_text)._p)
                                    logger.info("Added skill to end of skills section")
                    
                    elif section_type == 'projects':
                        for update in updates:
                            suggestion_text = update.get('suggestion', '')
                            project_name = update.get('role', '')  # Project name stored in 'role' field
                            
                            if suggestion_text:
                                if project_name:
                                    # Try to find matching project
                                    for k in range(i + 1, section_end):
                                        if project_name.lower() in doc.paragraphs[k].text.lower():
                                            next_para = k + 1
                                            while next_para < section_end and doc.paragraphs[next_para].text.strip().startswith('•'):
                                                next_para += 1
                                            doc.paragraphs[next_para - 1]._p.addnext(doc.add_paragraph(suggestion_text)._p)
                                            logger.info(f"Added suggestion under project: {project_name}")
                                            break
                                else:
                                    # Add to end of projects section if no specific project specified
                                    doc.paragraphs[section_end-1]._p.addnext(doc.add_paragraph(suggestion_text)._p)
                                    logger.info("Added project to end of section")
                    
                    elif section_type == 'experience':
                        logger.info("Processing experience section")
                        # Group suggestions by role
                        role_suggestions = {}
                        for update in updates:
                            role = update.get('role', '')
                            suggestion = update.get('suggestion', '')
                            if role and suggestion:
                                if role not in role_suggestions:
                                    role_suggestions[role] = []
                                role_suggestions[role].append(suggestion)
                        
                        # Process each role's suggestions
                        for role, suggestions in role_suggestions.items():
                            logger.info(f"Processing suggestions for role: {role}")
                            role_found = False
                            
                            # Find the specific role
                            for k in range(i + 1, section_end):
                                para_text = doc.paragraphs[k].text.strip()
                                if not para_text.startswith('•'):  # This is potentially a role line
                                    # More flexible role matching using key parts of the role
                                    role_key_parts = [part.lower() for part in role.split() if len(part) > 3]
                                    para_key_parts = [part.lower() for part in para_text.split() if len(part) > 3]
                                    
                                    # Check if enough key parts match
                                    matches = sum(1 for part in role_key_parts if any(part in p for p in para_key_parts))
                                    if matches >= min(2, len(role_key_parts)):  # At least 2 matches or all parts if less
                                        logger.info(f"Found matching role at index {k}: {para_text}")
                                        role_found = True
                                        
                                        # Find end of this role (start of next role or section end)
                                        role_end = section_end
                                        for j in range(k + 1, section_end):
                                            next_text = doc.paragraphs[j].text.strip()
                                            # Check for next role by looking for date patterns or known role titles
                                            if (not next_text.startswith('•') and 
                                                next_text and 
                                                j > k + 1 and
                                                (
                                                    # Date patterns
                                                    any(month in next_text for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']) or
                                                    # Common date formats
                                                    bool(re.search(r'\d{4}[-/]\d{4}|\d{4}[-/]Present|20\d{2}|19\d{2}', next_text)) or
                                                    # Company indicators
                                                    any(indicator in next_text for indicator in ['Inc.', 'LLC', 'Ltd.', 'Corp.', 'Corporation']) or
                                                    # Location patterns
                                                    bool(re.search(r'[A-Z][A-Za-z\s]+,\s*[A-Z]{2}', next_text)) or
                                                    # Role transition indicators
                                                    ':' in next_text or
                                                    'position' in next_text.lower() or
                                                    'role' in next_text.lower()
                                                )
                                            ):
                                                role_end = j
                                                break
                                        
                                        # Skip past company location
                                        start_idx = k
                                        for loc_idx in range(k + 1, role_end):
                                            if doc.paragraphs[loc_idx].text.strip() and not doc.paragraphs[loc_idx].text.strip().startswith('•'):
                                                start_idx = loc_idx
                                                break
                                        
                                        logger.info(f"Adding {len(suggestions)} suggestions under role between indices {start_idx} and {role_end}")
                                        add_experience_bullets(doc, start_idx, role_end, suggestions)
                                        break
                            
                            if not role_found:
                                logger.warning(f"Could not find matching role: {role}")

                    else:
                        # Handle any other section type by appending suggestions to the end
                        for update in updates:
                            suggestion_text = update.get('suggestion', '')
                            if suggestion_text:
                                doc.paragraphs[section_end-1]._p.addnext(doc.add_paragraph(suggestion_text)._p)
                                logger.info(f"Added suggestion to section: {section}")

        # Create downloads directory
        downloads_dir = os.path.join(os.getcwd(), 'downloads')
        os.makedirs(downloads_dir, exist_ok=True)

        # Save with timestamp
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        download_filename = f"updated_{timestamp}_{filename}"
        download_path = os.path.join(downloads_dir, download_filename)

        # Save the updated document
        doc.save(download_path)

        # return send_file(
        #     download_path,
        #     as_attachment=True,
        #     download_name=download_filename,
        #     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        # )

        return jsonify({'message': 'Resume updated successfully'}), 200
    except Exception as e:
        logger.error(f"Error in structured_resume_update: {str(e)}")
        return jsonify({'error': str(e)}), 500

    finally:
        # Cleanup temporary files
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            logger.error(f"Cleanup error: {e}")

def get_section_type(section_name: str) -> str:
    """Determine the type of section based on common variations of section names"""
    section_name = section_name.lower()
    
    # Map of section types to their common variations
    section_mappings = {
        'skills': [
            'skills', 'technical skills', 'core skills', 'competencies', 
            'expertise', 'qualifications', 'tools', 'technologies',
            'skills & tools', 'technical expertise'
        ],
        'projects': [
            'projects', 'project experience', 'technical projects',
            'relevant projects', 'key projects', 'professional projects',
            'computer science projects', 'software projects'
        ],
        'experience': [
            'experience', 'work experience', 'professional experience',
            'employment history', 'work history', 'career history',
            'relevant experience', 'professional background'
        ]
    }
    
    # Check each section type
    for section_type, variations in section_mappings.items():
        if any(variation in section_name for variation in variations):
            return section_type
            
    return 'other'

def add_experience_bullets(doc, role_start_idx: int, role_end_idx: int, suggestions: List[str]) -> None:
    """Add bullet points with consistent formatting under a role"""
    logger.info(f"Adding {len(suggestions)} bullet points between indices {role_start_idx} and {role_end_idx}")
    
    # Find existing bullet points to use as template for formatting
    template_bullet = None
    for i in range(role_start_idx + 1, role_end_idx):
        if doc.paragraphs[i].text.strip().startswith('•'):
            template_bullet = doc.paragraphs[i]
            break
    
    # Find all existing bullet points under this role
    bullet_indices = []
    for i in range(role_start_idx + 1, role_end_idx):
        if doc.paragraphs[i].text.strip().startswith('•'):
            bullet_indices.append(i)
    
    logger.info(f"Found {len(bullet_indices)} existing bullet points")
    
    # If no existing bullets, add all suggestions after role with default formatting
    if not bullet_indices:
        for suggestion in suggestions:
            new_para = doc.add_paragraph(suggestion)
            new_para.paragraph_format.left_indent = Inches(0.5)  # Default indentation
            doc.paragraphs[role_start_idx]._p.addnext(new_para._p)
            logger.info(f"Added bullet point after role: {suggestion}")
        return

    # Distribute suggestions across existing bullets
    for idx, suggestion in enumerate(suggestions):
        # Choose insertion point based on position in suggestions list
        position = (idx * len(bullet_indices)) // len(suggestions)
        insert_idx = bullet_indices[position]
        
        # Create new paragraph with matching formatting
        new_para = doc.add_paragraph()
        
        # Copy formatting from template if available
        if template_bullet:
            # Copy paragraph format (indentation, spacing, etc.)
            if template_bullet._element.pPr is not None:
                new_para._element.get_or_add_pPr().append(
                    deepcopy(template_bullet._element.pPr)
                )
            
            # Copy run format (font, size, etc.)
            if template_bullet.runs:
                new_run = new_para.add_run(suggestion)
                template_run = template_bullet.runs[0]
                if hasattr(template_run.font, 'name'):
                    new_run.font.name = template_run.font.name
                if hasattr(template_run.font, 'size'):
                    new_run.font.size = template_run.font.size
                new_run.font.bold = template_run.font.bold
                new_run.font.italic = template_run.font.italic
            else:
                new_para.add_run(suggestion)
        else:
            # Default formatting if no template
            new_para.add_run(suggestion)
            new_para.paragraph_format.left_indent = Inches(0.5)
        
        # Insert at the chosen position
        doc.paragraphs[insert_idx]._p.addnext(new_para._p)
        logger.info(f"Added formatted bullet point at position {position}: {suggestion}")

if __name__ == '__main__':
    # Create UPLOAD_FOLDER if it doesn't exist
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    app.run(debug=True)

