#!/usr/bin/env python3
"""Debug script to examine DOCX structure and understand mapping issues."""

import sys
from docx import Document

def debug_docx_structure():
    """Examine the DOCX structure to understand mapping issues."""
    
    # Load the document
    doc = Document('nagarjuna_kocharla_resume_main.docx')
    
    print('=== DOCX DOCUMENT STRUCTURE ===')
    print(f'Total paragraphs: {len(doc.paragraphs)}')
    
    # Show first 30 paragraphs to see structure
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        if text:
            print(f'{i:2d}: {text[:80]}{"..." if len(text) > 80 else ""}')
    
    print('\n=== BULLET DETECTION ===')
    bullet_count = 0
    bullet_positions = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('•') or text.startswith('-') or text.startswith('*'):
            bullet_count += 1
            bullet_positions.append(i)
            print(f'Bullet {bullet_count} at paragraph {i}: {text[:80]}...')
    
    print(f'\nTotal bullets found: {bullet_count}')
    print(f'Bullet positions: {bullet_positions}')
    
    print('\n=== ROLE DETECTION ===')
    role_count = 0
    role_positions = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if ' at ' in text and any(keyword in text.lower() for keyword in ['inc', 'corp', 'llc', 'ltd', 'company', 'tech', 'systems', 'consultancy', 'international']):
            role_count += 1
            role_positions.append(i)
            print(f'Role {role_count} at paragraph {i}: {text}')
    
    print(f'\nTotal roles found: {role_count}')
    print(f'Role positions: {role_positions}')
    
    print('\n=== MAPPING ANALYSIS ===')
    print('The issue is likely:')
    print('1. Resume Parser extracts bullets from text with IDs: bullet_1, bullet_2, etc.')
    print('2. Document Editor tries to find these IDs in the DOCX structure')
    print('3. But the DOCX paragraphs may not match exactly due to formatting differences')
    print('4. Text similarity threshold (0.7) may be too strict for real resumes')

if __name__ == "__main__":
    debug_docx_structure()
